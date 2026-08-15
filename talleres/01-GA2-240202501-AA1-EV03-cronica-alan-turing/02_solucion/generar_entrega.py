from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TALLER_DIR = ROOT.parent
OUTPUT_DIR = TALLER_DIR / "03_entrega"
OUTPUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUTPUT_DIR / "GA2-240202501-AA1-EV03_Cronica_Alan_Turing.docx"
TEXT_PATH = OUTPUT_DIR / "GA2-240202501-AA1-EV03_Cronica_Alan_Turing.txt"

GREEN = "39A900"
DARK_GREEN = "174C2C"
PALE_GREEN = "EDF7E9"
INK = "18221B"
GRAY = "5E6A62"
LIGHT_GRAY = "F4F6F4"
MID_GRAY = "D8DED9"
WHITE = "FFFFFF"


LEAD = (
    "In September 1939, while Europe entered the Second World War, a quiet young "
    "mathematician arrived at Bletchley Park, Britain’s secret codebreaking center. "
    "His name was Alan Turing, and he faced a problem that seemed impossible: German "
    "Enigma messages changed their settings every day. His response to that challenge "
    "changed the history of computing."
)

SECTIONS = [
    (
        "1912–1936  |  A CURIOUS AND INDEPENDENT MIND",
        [
            (
                "Alan Mathison Turing was born in London on June 23, 1912. As a boy, "
                "he showed a strong interest in numbers, science, and experiments. Later, "
                "he studied mathematics at King’s College, Cambridge. In formal photographs, "
                "he was clean-shaven, had short dark hair, and wore a simple suit and tie. "
                "He was tall and slim, and he often looked serious. People described him as "
                "shy and direct, but he was also curious, imaginative, humorous, and persistent. "
                "He enjoyed cycling and became a dedicated long-distance runner."
            ),
            (
                "In 1936, Turing published a paper about an abstract machine. The machine "
                "read symbols, followed precise rules, and solved problems one step at a time. "
                "His idea showed that one machine could perform many different tasks when it "
                "received different instructions. This concept became one of the foundations "
                "of theoretical computer science and programming."
            ),
        ],
    ),
    (
        "1939–1945  |  BREAKING IMPOSSIBLE CODES",
        [
            (
                "At Bletchley Park, Turing led Hut 8’s work on German naval Enigma. He did "
                "not solve the challenge alone. He collaborated with mathematicians, engineers, "
                "linguists, operators, and other specialists. Building on earlier Polish "
                "codebreaking work, Turing, Gordon Welchman, and the wider British team designed "
                "the Bombe. This electromechanical machine checked many possible Enigma settings. "
                "Turing also developed a statistical method called Banburismus. Their careful "
                "analysis helped the Allies read important messages about German submarines."
            ),
            (
                "Life at the secret center was intense. Turing worked for long hours, "
                "concentrated deeply, and questioned old methods. He sometimes behaved in an "
                "unconventional way; he even chained his tea mug to a radiator so that nobody "
                "used it. However, his unusual habits did not stop teamwork. When one idea failed, "
                "he tested another. In my view, his persistence, logical thinking, and creativity "
                "resembled the qualities of a good software analyst. He understood the problem, "
                "designed a process, tested results, and improved the solution."
            ),
        ],
    ),
    (
        "1945–1950  |  FROM THEORY TO SOFTWARE",
        [
            (
                "After the war, Turing joined the National Physical Laboratory. In 1945, he "
                "designed the Automatic Computing Engine, or ACE. It was an electronic computer "
                "design that stored instructions and data in memory. Progress was slow, so in "
                "1948 he moved to the University of Manchester. There, he developed routines and "
                "documentation for early computers. His work connected mathematical ideas with "
                "practical programming."
            ),
            (
                "In 1950, Turing published “Computing Machinery and Intelligence.” He asked "
                "whether machines could think and described the imitation game, which later "
                "became known as the Turing Test. Once again, he transformed a complex question "
                "into a clear procedure that people could discuss and test."
            ),
        ],
    ),
    (
        "1952–1954  |  INJUSTICE AND LOSS",
        [
            (
                "His final years brought a painful injustice. In 1952, British authorities "
                "prosecuted Turing because he had a relationship with another man, which was "
                "illegal in Britain at that time. He accepted hormonal treatment instead of "
                "going to prison. On June 7, 1954, he died from cyanide poisoning at the age of "
                "41; the inquest recorded a verdict of suicide. Decades later, the British "
                "government apologized for his treatment, and Queen Elizabeth II granted him a "
                "royal pardon in 2013."
            ),
        ],
    ),
    (
        "WHY HIS STORY MATTERS TO SOFTWARE DEVELOPERS",
        [
            (
                "Turing’s legacy remains important today. His ideas influenced algorithms, "
                "programmable computers, cybersecurity, and artificial intelligence. For students "
                "of Software Analysis and Development, his method offers a clear lesson: define "
                "the problem, divide it into smaller tasks, write precise instructions, test possible "
                "answers, and collaborate with a team. I believe Alan Turing was an iconic pioneer "
                "because he combined curiosity with disciplined work and used logical ideas to solve "
                "real problems. His achievements changed technology, while his life also showed why "
                "society must defend dignity, equality, and respect."
            ),
        ],
    ),
]

SOURCES = [
    (
        "Bletchley Park Trust. (n.d.). Alan Turing.",
        "https://bletchleypark.org.uk/wp-content/uploads/record_attachments/1800.pdf",
    ),
    (
        "Bank of England. (2019). Alan Turing to be the face of the new £50 note.",
        "https://www.bankofengland.co.uk/-/media/boe/files/news/2019/july/alan-turing-to-be-the-face-of-new-50-note.pdf",
    ),
    (
        "Computer History Museum. (n.d.). Pilot ACE.",
        "https://www.computerhistory.org/revolution/story/96",
    ),
]


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size=6, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    edges = ["top", "left", "bottom", "right"]
    if inside:
        edges.extend(["insideH", "insideV"])
    for edge in edges:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_paragraph_border(paragraph, color=GREEN, size=18, space=8, side="left"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def set_keep_together(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    p_pr.append(keep_lines)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE  ")
    run.font.name = "Lato"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "2"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, value, end):
        run._r.append(node)


def add_hyperlink(paragraph, text, url, color=DARK_GREEN, underline=True):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    run_properties.append(color_element)
    if underline:
        underline_element = OxmlElement("w:u")
        underline_element.set(qn("w:val"), "single")
        run_properties.append(underline_element)
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Lato")
    font.set(qn("w:hAnsi"), "Lato")
    run_properties.append(font)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    run_properties.append(size)
    new_run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_run_font(run, size=None, bold=None, color=None, italic=None, caps=None):
    run.font.name = "Lato"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.font.italic = italic
    if caps is not None:
        run.font.all_caps = caps


def add_section_heading(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    set_keep_with_next(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=True, color=DARK_GREEN, caps=True)
    set_paragraph_border(paragraph, color=GREEN, size=16, space=6, side="bottom")
    return paragraph


def add_body_paragraph(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.22)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.08
    set_keep_together(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.25, color=INK)
    return paragraph


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True

    properties = document.core_properties
    properties.title = "Alan Turing: The Quiet Mind Behind Modern Computing"
    properties.subject = "Crónica en inglés — GA2-240202501-AA1-EV03"
    properties.author = "Jhon Steven Alvarez Ruiz"
    properties.keywords = "SENA, English, chronicle, simple past, Alan Turing, software"
    properties.comments = "Documento elaborado con base en la rúbrica IE-GA2-240202501-AA1-EV03."

    normal = document.styles["Normal"]
    normal.font.name = "Lato"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Lato")
    normal.font.size = Pt(10.25)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)

    # Header on content pages.
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(7.06))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(1.0)
    header_table.columns[1].width = Inches(6.06)
    remove_table_borders(header_table)
    left_cell, right_cell = header_table.rows[0].cells
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(left_cell, top=0, start=0, bottom=0, end=0)
    set_cell_margins(right_cell, top=0, start=0, bottom=0, end=0)
    header_brand = left_cell.paragraphs[0]
    header_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_brand_run = header_brand.add_run("SENA")
    set_run_font(header_brand_run, size=11, bold=True, color=GREEN, caps=True)
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_p.add_run("ENGLISH CHRONICLE  ·  GA2-240202501-AA1-EV03")
    set_run_font(right_run, size=7.5, bold=True, color=GRAY, caps=True)
    set_paragraph_border(right_p, color=GREEN, size=10, space=4, side="bottom")

    # Footer on content pages.
    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(7.06))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.columns[0].width = Inches(5.7)
    footer_table.columns[1].width = Inches(1.36)
    remove_table_borders(footer_table)
    footer_left, footer_right = footer_table.rows[0].cells
    set_cell_margins(footer_left, top=0, start=0, bottom=0, end=0)
    set_cell_margins(footer_right, top=0, start=0, bottom=0, end=0)
    footer_text = footer_left.paragraphs[0]
    footer_text_run = footer_text.add_run("ANÁLISIS Y DESARROLLO DE SOFTWARE")
    set_run_font(footer_text_run, size=7.5, bold=True, color=GRAY, caps=True)
    add_page_number(footer_right.paragraphs[0])

    # Cover.
    cover_brand = document.add_table(rows=1, cols=1)
    cover_brand.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_brand.autofit = False
    remove_table_borders(cover_brand)
    brand_cell = cover_brand.cell(0, 0)
    brand_cell.width = Inches(1.45)
    set_cell_shading(brand_cell, GREEN)
    set_cell_margins(brand_cell, top=100, start=120, bottom=100, end=120)
    brand_paragraph = brand_cell.paragraphs[0]
    brand_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand_paragraph.add_run("SENA")
    set_run_font(brand_run, size=21, bold=True, color=WHITE, caps=True)

    cover_spacer = document.add_paragraph()
    cover_spacer.paragraph_format.space_after = Pt(0)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(8)
    kicker_run = kicker.add_run("EVIDENCIA DE PRODUCTO  ·  CRÓNICA EN INGLÉS")
    set_run_font(kicker_run, size=9, bold=True, color=DARK_GREEN, caps=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    title_run = title.add_run("ALAN TURING")
    set_run_font(title_run, size=29, bold=True, color=INK, caps=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(5)
    subtitle_run = subtitle.add_run("The Quiet Mind Behind Modern Computing")
    set_run_font(subtitle_run, size=16, bold=True, color=DARK_GREEN)

    deck = document.add_paragraph()
    deck.alignment = WD_ALIGN_PARAGRAPH.CENTER
    deck.paragraph_format.space_after = Pt(18)
    deck_run = deck.add_run("A chronicle of curiosity, codes, and courage")
    set_run_font(deck_run, size=10.5, italic=True, color=GRAY)

    rule_table = document.add_table(rows=1, cols=3)
    rule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(rule_table)
    widths = [1.55, 3.95, 1.55]
    colors = [PALE_GREEN, GREEN, PALE_GREEN]
    for cell, width, fill in zip(rule_table.rows[0].cells, widths, colors):
        cell.width = Inches(width)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=45, start=0, bottom=45, end=0)

    info = document.add_table(rows=1, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    info.columns[0].width = Inches(2.0)
    info.columns[1].width = Inches(4.55)
    set_table_borders(info, color=MID_GRAY, size=6, inside=True)
    info_data = [("Aprendiz", "Jhon Steven Alvarez Ruiz")]
    for row_index, (label, value) in enumerate(info_data):
        label_cell, value_cell = info.rows[row_index].cells
        label_cell.width = Inches(2.0)
        value_cell.width = Inches(4.55)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(label_cell, PALE_GREEN)
        set_cell_shading(value_cell, WHITE if row_index % 2 == 0 else LIGHT_GRAY)
        set_cell_margins(label_cell, top=105, start=140, bottom=105, end=120)
        set_cell_margins(value_cell, top=105, start=140, bottom=105, end=120)
        label_p = label_cell.paragraphs[0]
        value_p = value_cell.paragraphs[0]
        label_run = label_p.add_run(label)
        value_run = value_p.add_run(value)
        set_run_font(label_run, size=9, bold=True, color=DARK_GREEN)
        set_run_font(value_run, size=9.5, color=INK)

    cover_note = document.add_paragraph()
    cover_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_note.paragraph_format.space_before = Pt(18)
    note_run = cover_note.add_run("SERVICIO NACIONAL DE APRENDIZAJE · SENA")
    set_run_font(note_run, size=8, bold=True, color=GRAY, caps=True)

    document.add_page_break()

    # Chronicle title and lead.
    content_kicker = document.add_paragraph()
    content_kicker.paragraph_format.space_after = Pt(2)
    content_kicker_run = content_kicker.add_run("ENGLISH CHRONICLE")
    set_run_font(content_kicker_run, size=8.5, bold=True, color=GREEN, caps=True)

    content_title = document.add_paragraph()
    content_title.paragraph_format.space_after = Pt(1)
    set_keep_with_next(content_title)
    content_title_run = content_title.add_run("Alan Turing: The Quiet Mind\nBehind Modern Computing")
    set_run_font(content_title_run, size=23, bold=True, color=INK)

    content_deck = document.add_paragraph()
    content_deck.paragraph_format.space_after = Pt(8)
    content_deck_run = content_deck.add_run("A chronicle of curiosity, codes, and courage")
    set_run_font(content_deck_run, size=10.5, italic=True, color=GRAY)

    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead.paragraph_format.left_indent = Inches(0.12)
    lead.paragraph_format.right_indent = Inches(0.08)
    lead.paragraph_format.space_before = Pt(2)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.06
    set_paragraph_border(lead, color=GREEN, size=26, space=8, side="left")
    set_paragraph_shading(lead, PALE_GREEN)
    lead_run = lead.add_run(LEAD)
    set_run_font(lead_run, size=10.5, bold=True, color=INK)

    # Compact timeline supporting the chronological structure.
    timeline = document.add_table(rows=2, cols=6)
    timeline.alignment = WD_TABLE_ALIGNMENT.CENTER
    timeline.autofit = False
    remove_table_borders(timeline)
    timeline_data = [
        ("1912", "Born"),
        ("1936", "Machine idea"),
        ("1939", "Bletchley Park"),
        ("1945", "ACE design"),
        ("1950", "Turing Test"),
        ("1954", "Final year"),
    ]
    for column, (year, label) in enumerate(timeline_data):
        year_cell = timeline.cell(0, column)
        label_cell = timeline.cell(1, column)
        year_cell.width = Inches(1.08)
        label_cell.width = Inches(1.08)
        set_cell_shading(year_cell, DARK_GREEN if column % 2 == 0 else GREEN)
        set_cell_shading(label_cell, LIGHT_GRAY)
        set_cell_margins(year_cell, top=40, start=40, bottom=40, end=40)
        set_cell_margins(label_cell, top=50, start=30, bottom=50, end=30)
        year_p = year_cell.paragraphs[0]
        label_p = label_cell.paragraphs[0]
        year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year_run = year_p.add_run(year)
        label_run = label_p.add_run(label)
        set_run_font(year_run, size=8.5, bold=True, color=WHITE)
        set_run_font(label_run, size=7.5, bold=True, color=GRAY)

    document.add_paragraph().paragraph_format.space_after = Pt(0)

    # First two eras on page 2.
    for section_title, paragraphs in SECTIONS[:2]:
        add_section_heading(document, section_title)
        for paragraph_text in paragraphs:
            add_body_paragraph(document, paragraph_text)

    # Intentional page break keeps headings and sources balanced.
    document.add_page_break()

    for section_title, paragraphs in SECTIONS[2:4]:
        add_section_heading(document, section_title)
        for paragraph_text in paragraphs:
            add_body_paragraph(document, paragraph_text)

    method_table = document.add_table(rows=1, cols=4)
    method_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    method_table.autofit = False
    remove_table_borders(method_table)
    for index, label in enumerate(("ANALYZE", "DESIGN", "TEST", "IMPROVE")):
        cell = method_table.cell(0, index)
        cell.width = Inches(1.62)
        set_cell_shading(cell, DARK_GREEN if index in (0, 3) else GREEN)
        set_cell_margins(cell, top=80, start=40, bottom=80, end=40)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        set_run_font(run, size=8.5, bold=True, color=WHITE, caps=True)

    final_title, final_paragraphs = SECTIONS[4]
    add_section_heading(document, final_title)
    for paragraph_text in final_paragraphs:
        add_body_paragraph(document, paragraph_text)

    sources_heading = document.add_paragraph()
    sources_heading.paragraph_format.space_before = Pt(6)
    sources_heading.paragraph_format.space_after = Pt(2)
    set_keep_with_next(sources_heading)
    sources_heading_run = sources_heading.add_run("SOURCES CONSULTED")
    set_run_font(sources_heading_run, size=8, bold=True, color=GRAY, caps=True)

    for source_text, source_url in SOURCES:
        paragraph = document.add_paragraph(style=None)
        paragraph.paragraph_format.left_indent = Inches(0.16)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        paragraph.paragraph_format.space_after = Pt(1.5)
        bullet_run = paragraph.add_run("• ")
        set_run_font(bullet_run, size=8.5, bold=True, color=GREEN)
        add_hyperlink(paragraph, source_text, source_url, color=DARK_GREEN, underline=False)

    # Explicit language tag improves spell-check behavior in Word/LibreOffice.
    body = document._body._element
    for run_properties in body.iter(qn("w:rPr")):
        lang = run_properties.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            run_properties.append(lang)
        lang.set(qn("w:val"), "en-US")

    document.save(DOCX_PATH)

    all_paragraphs = [LEAD]
    for _, paragraphs in SECTIONS:
        all_paragraphs.extend(paragraphs)
    word_count = len(re.findall(r"\b[\w’'-]+\b", " ".join(all_paragraphs), re.UNICODE))

    text_lines = [
        "ALAN TURING: THE QUIET MIND BEHIND MODERN COMPUTING",
        "A chronicle of curiosity, codes, and courage",
        "",
        LEAD,
        "",
    ]
    for section_title, paragraphs in SECTIONS:
        text_lines.append(section_title)
        text_lines.extend(paragraphs)
        text_lines.append("")
    text_lines.append("SOURCES CONSULTED")
    for source_text, source_url in SOURCES:
        text_lines.append(f"- {source_text} {source_url}")
    text_lines.extend(["", f"Chronicle word count: {word_count}"])
    TEXT_PATH.write_text("\n\n".join(text_lines), encoding="utf-8")
    return word_count


if __name__ == "__main__":
    count = build_document()
    print(f"Created: {DOCX_PATH}")
    print(f"Created: {TEXT_PATH}")
    print(f"Chronicle word count: {count}")
