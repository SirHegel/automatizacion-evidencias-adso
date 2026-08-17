"""Genera el correo de postulacion laboral en DOCX y texto plano."""

from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOLUTION_DIR = Path(__file__).resolve().parent
WORKSHOP_DIR = SOLUTION_DIR.parent
DELIVERY_DIR = WORKSHOP_DIR / "03_entrega"
SOURCE_PATH = SOLUTION_DIR / "TEXTO_CORREO.txt"
DOCX_PATH = DELIVERY_DIR / "GA2-240202501-AA2-EV03_Correo_Solicitud_Empleo.docx"
TEXT_PATH = DELIVERY_DIR / "GA2-240202501-AA2-EV03_Correo_Solicitud_Empleo.txt"

GREEN = "39A900"
DARK_GREEN = "174C2C"
PALE_GREEN = "EDF7E9"
INK = "18221B"
MUTED = "5E6A62"
LINE = "D8DED9"
WHITE = "FFFFFF"


def parse_source():
    text = SOURCE_PATH.read_text(encoding="utf-8").strip()
    blocks = text.split("\n\n")
    header_lines = blocks[0].splitlines()
    if len(header_lines) != 3:
        raise ValueError("El texto fuente debe comenzar con From, To y Subject.")

    headers = {}
    for expected, line in zip(("From", "To", "Subject"), header_lines, strict=True):
        prefix = f"{expected}: "
        if not line.startswith(prefix):
            raise ValueError(f"Falta el encabezado {expected} en el texto fuente.")
        headers[expected] = line.removeprefix(prefix)

    message_blocks = blocks[1:]
    words = re.findall(r"[A-Za-z]+(?:['’][A-Za-z]+)?", " ".join(message_blocks))
    if not 200 <= len(words) <= 400:
        raise ValueError(
            f"El mensaje contiene {len(words)} palabras; la guia exige entre 200 y 400."
        )
    return text + "\n", headers, message_blocks, len(words)


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=50, start=130, bottom=50, end=130):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6):
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def keep_together(paragraph):
    properties = paragraph._p.get_or_add_pPr()
    properties.append(OxmlElement("w:keepLines"))


def keep_with_next(paragraph):
    properties = paragraph._p.get_or_add_pPr()
    properties.append(OxmlElement("w:keepNext"))


def format_paragraph(paragraph, *, after=0, before=0, alignment=None):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    if alignment is not None:
        paragraph.alignment = alignment
    keep_together(paragraph)


def format_run(run, *, bold=False, color=INK):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, *, bold=False, color=INK):
    run = paragraph.add_run(text)
    format_run(run, bold=bold, color=color)
    return run


def build_document(headers, message_blocks, word_count):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)

    properties = document.core_properties
    properties.title = "Application for the Junior Software Developer Position"
    properties.subject = "GA2-240202501-AA2-EV03 — Written job application email"
    properties.author = "Jhon Steven Alvarez Ruiz"
    properties.last_modified_by = "Jhon Steven Alvarez Ruiz"
    properties.keywords = "SENA, English, job application, software development"

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)

    banner = document.add_table(rows=1, cols=2)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.columns[0].width = Inches(3.6)
    banner.columns[1].width = Inches(3.45)
    for cell in banner.rows[0].cells:
        set_cell_shading(cell, DARK_GREEN)
        set_cell_margins(cell, top=45, bottom=45)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left = banner.cell(0, 0).paragraphs[0]
    format_paragraph(left, after=0)
    add_text(left, "SENA · ENGLISH EVIDENCE", bold=True, color=WHITE)
    right = banner.cell(0, 1).paragraphs[0]
    format_paragraph(right, after=0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_text(right, "GA2-240202501-AA2-EV03", bold=True, color=WHITE)

    title = document.add_paragraph()
    format_paragraph(title, before=2, after=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    keep_with_next(title)
    add_text(title, "JOB APPLICATION EMAIL", bold=True, color=DARK_GREEN)

    metadata = document.add_table(rows=3, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata.autofit = False
    metadata.columns[0].width = Inches(1.15)
    metadata.columns[1].width = Inches(5.9)
    set_table_borders(metadata)
    for row_index, label in enumerate(("FROM", "TO", "SUBJECT")):
        label_cell, value_cell = metadata.rows[row_index].cells
        label_cell.width = Inches(1.15)
        value_cell.width = Inches(5.9)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(label_cell, PALE_GREEN)
        set_cell_margins(label_cell)
        set_cell_margins(value_cell)
        label_paragraph = label_cell.paragraphs[0]
        value_paragraph = value_cell.paragraphs[0]
        format_paragraph(label_paragraph, after=0)
        format_paragraph(value_paragraph, after=0)
        add_text(label_paragraph, label, bold=True, color=DARK_GREEN)
        add_text(value_paragraph, headers[label.title()], bold=label == "SUBJECT")
    set_repeat_table_header(metadata.rows[0])

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(2)
    spacer.add_run("")

    for index, block in enumerate(message_blocks):
        paragraph = document.add_paragraph()
        is_signature = index == len(message_blocks) - 1
        is_salutation = index == 0
        is_closing = block == "Sincerely,"
        format_paragraph(paragraph, after=0 if is_signature or is_closing else 3)
        if is_salutation or is_closing:
            keep_with_next(paragraph)
        lines = block.splitlines()
        for line_index, line in enumerate(lines):
            add_text(
                paragraph,
                line,
                bold=is_salutation or is_closing or (is_signature and line_index == 0),
            )
            if line_index < len(lines) - 1:
                paragraph.add_run().add_break()

    document.core_properties.comments = (
        f"Mensaje de {word_count} palabras; Arial 12; interlineado 1,5."
    )
    document.save(DOCX_PATH)


def main():
    source_text, headers, message_blocks, word_count = parse_source()
    DELIVERY_DIR.mkdir(parents=True, exist_ok=True)
    TEXT_PATH.write_text(source_text, encoding="utf-8")
    build_document(headers, message_blocks, word_count)
    print(f"Created: {DOCX_PATH}")
    print(f"Created: {TEXT_PATH}")
    print(f"Message word count: {word_count}")


if __name__ == "__main__":
    main()
