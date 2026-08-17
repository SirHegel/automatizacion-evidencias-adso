"""Genera la evidencia sobre algoritmos de edad y año bisiesto."""

from __future__ import annotations

from datetime import date
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOLUTION_DIR = Path(__file__).resolve().parent
WORKSHOP_DIR = SOLUTION_DIR.parent
DELIVERY_DIR = WORKSHOP_DIR / "03_entrega"
PSEUDOCODE_DIR = SOLUTION_DIR / "pseudocodigo"
DIAGRAM_DIR = SOLUTION_DIR / "recursos" / "diagramas"

AGE_SOURCE = PSEUDOCODE_DIR / "01_calcular_edad.psc"
LEAP_SOURCE = PSEUDOCODE_DIR / "02_determinar_anio_bisiesto.psc"
AGE_DIAGRAM = DIAGRAM_DIR / "01_calcular_edad.png"
LEAP_DIAGRAM = DIAGRAM_DIR / "02_determinar_anio_bisiesto.png"
DOCX_PATH = DELIVERY_DIR / "GA3-220501093-AA2-EV01_Fundamentos_Programacion_Estructurada.docx"

GREEN = "39A900"
DARK_GREEN = "174C2C"
DEEP_GREEN = "0E3322"
PALE_GREEN = "EDF7E9"
PALE_BLUE = "EAF3F7"
PALE_GOLD = "FFF4D6"
PALE_RED = "FCECE8"
INK = "18221B"
MUTED = "5E6A62"
LINE = "D8DED9"
WHITE = "FFFFFF"
CODE_BG = "F4F6F4"

FONT_REGULAR_CANDIDATES = (
    Path("/usr/share/fonts/truetype/msttcorefonts/Arial.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
)
FONT_BOLD_CANDIDATES = (
    Path("/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
)


AGE_TESTS = (
    ("Cumpleaños pendiente", (2000, 8, 20), (2026, 8, 16), "25"),
    ("Cumpleaños exacto", (2000, 8, 16), (2026, 8, 16), "26"),
    ("Cumpleaños cumplido", (2000, 8, 15), (2026, 8, 16), "26"),
    ("Nacimiento en día bisiesto", (2004, 2, 29), (2025, 2, 28), "20"),
    ("Aniversario tras febrero bisiesto", (2004, 2, 29), (2025, 3, 1), "21"),
    ("Fecha inexistente", (2023, 2, 29), (2026, 8, 16), "ERROR"),
    ("Nacimiento posterior", (2027, 1, 1), (2026, 8, 16), "ERROR"),
)

LEAP_TESTS = (
    (2000, "Sí", "Divisible por 400"),
    (1900, "No", "Divisible por 100, pero no por 400"),
    (2024, "Sí", "Divisible por 4 y no por 100"),
    (2023, "No", "No es divisible por 4"),
    (0, "ERROR", "El año debe ser positivo"),
)


def is_leap_year(year: int) -> bool:
    if year <= 0:
        raise ValueError("El año debe ser mayor que cero.")
    return year % 400 == 0 or (year % 4 == 0 and year % 100 != 0)


def valid_date(value: tuple[int, int, int]) -> bool:
    try:
        date(*value)
    except ValueError:
        return False
    return True


def calculate_age(
    birth: tuple[int, int, int],
    current: tuple[int, int, int],
) -> int:
    if not valid_date(birth) or not valid_date(current):
        raise ValueError("Las fechas deben existir en el calendario gregoriano.")
    birth_date = date(*birth)
    current_date = date(*current)
    if birth_date > current_date:
        raise ValueError("La fecha de nacimiento no puede ser posterior a la actual.")
    age = current_date.year - birth_date.year
    if (current_date.month, current_date.day) < (birth_date.month, birth_date.day):
        age -= 1
    return age


def run_tests() -> None:
    for name, birth, current, expected in AGE_TESTS:
        try:
            result = str(calculate_age(birth, current))
        except ValueError:
            result = "ERROR"
        if result != expected:
            raise AssertionError(f"Falló la prueba de edad '{name}': {result} != {expected}")

    for year, expected, _reason in LEAP_TESTS:
        try:
            result = "Sí" if is_leap_year(year) else "No"
        except ValueError:
            result = "ERROR"
        if result != expected:
            raise AssertionError(f"Falló la prueba de bisiesto {year}: {result} != {expected}")


def first_existing(candidates: tuple[Path, ...]) -> Path:
    for path in candidates:
        if path.is_file():
            return path
    raise FileNotFoundError(f"No se encontró ninguna tipografía: {candidates}")


def image_font(size: int, *, bold: bool = False):
    candidates = FONT_BOLD_CANDIDATES if bold else FONT_REGULAR_CANDIDATES
    return ImageFont.truetype(str(first_existing(candidates)), size=size)


def wrap_for_width(draw, text: str, font, max_width: int) -> str:
    lines: list[str] = []
    for original_line in text.splitlines() or [""]:
        words = original_line.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return "\n".join(lines)


def draw_centered_text(draw, bounds, text, *, font, fill="#18221B", padding=30):
    left, top, right, bottom = bounds
    wrapped = wrap_for_width(draw, text, font, right - left - 2 * padding)
    box = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=8, align="center")
    width = box[2] - box[0]
    height = box[3] - box[1]
    x = (left + right - width) / 2
    y = (top + bottom - height) / 2 - box[1]
    draw.multiline_text((x, y), wrapped, font=font, fill=fill, spacing=8, align="center")


def draw_node(draw, kind: str, bounds, text: str, *, fill: str, outline: str = "#174C2C"):
    left, top, right, bottom = bounds
    if kind == "terminator":
        draw.rounded_rectangle(bounds, radius=(bottom - top) // 2, fill=fill, outline=outline, width=5)
    elif kind == "input":
        offset = min(55, (right - left) // 8)
        points = [(left + offset, top), (right, top), (right - offset, bottom), (left, bottom)]
        draw.polygon(points, fill=fill)
        draw.line(points + [points[0]], fill=outline, width=5, joint="curve")
    elif kind == "decision":
        points = [((left + right) // 2, top), (right, (top + bottom) // 2), ((left + right) // 2, bottom), (left, (top + bottom) // 2)]
        draw.polygon(points, fill=fill)
        draw.line(points + [points[0]], fill=outline, width=5, joint="curve")
    else:
        draw.rounded_rectangle(bounds, radius=18, fill=fill, outline=outline, width=5)
    draw_centered_text(draw, bounds, text, font=image_font(34, bold=True), padding=55)


def draw_arrow(draw, points, *, label: str | None = None, label_at=None, color="#174C2C"):
    draw.line(points, fill=color, width=6, joint="curve")
    start = points[-2]
    end = points[-1]
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 24
    spread = math.pi / 6
    arrow = [
        end,
        (
            end[0] - length * math.cos(angle - spread),
            end[1] - length * math.sin(angle - spread),
        ),
        (
            end[0] - length * math.cos(angle + spread),
            end[1] - length * math.sin(angle + spread),
        ),
    ]
    draw.polygon(arrow, fill=color)
    if label and label_at:
        font = image_font(28, bold=True)
        box = draw.textbbox((0, 0), label, font=font)
        width = box[2] - box[0]
        height = box[3] - box[1]
        x, y = label_at
        draw.rounded_rectangle(
            (x - 12, y - 8, x + width + 12, y + height + 8),
            radius=10,
            fill="#FFFFFF",
            outline="#D8DED9",
            width=2,
        )
        draw.text((x, y), label, font=font, fill=color)


def diagram_base(title: str, subtitle: str, *, height: int = 2300):
    image = Image.new("RGB", (1700, height), "#F7FAF7")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((55, 35, 1645, 130), radius=22, fill="#174C2C")
    draw.text((90, 55), title, font=image_font(42, bold=True), fill="#FFFFFF")
    subtitle_font = image_font(26)
    draw.text((85, 145), subtitle, font=subtitle_font, fill="#5E6A62")
    return image, draw


def build_age_diagram(path: Path) -> None:
    image, draw = diagram_base(
        "ALGORITMO 1 · EDAD ACTUAL",
        "Validación, decisión de cumpleaños y repetición controlada",
    )
    start = (610, 210, 1090, 305)
    input_dates = (350, 375, 1350, 555)
    valid = (500, 650, 1200, 900)
    error = (1250, 680, 1640, 870)
    calculate = (430, 995, 1270, 1135)
    birthday = (510, 1230, 1190, 1480)
    adjust = (1250, 1270, 1630, 1435)
    output = (500, 1580, 1200, 1715)
    again = (560, 1810, 1140, 2030)
    end = (650, 2160, 1050, 2250)

    draw_node(draw, "terminator", start, "INICIO", fill="#EAF3F7")
    draw_node(draw, "input", input_dates, "Ingresar fecha de nacimiento\ny fecha actual", fill="#EAF3F7")
    draw_node(draw, "decision", valid, "¿Fechas gregorianas válidas\ny nacimiento no posterior?", fill="#FFF4D6")
    draw_node(draw, "input", error, "Mostrar error", fill="#FCECE8", outline="#A84232")
    draw_node(draw, "process", calculate, "edad ← año actual − año de nacimiento", fill="#EDF7E9")
    draw_node(draw, "decision", birthday, "¿La fecha actual es anterior\nal cumpleaños de ese año?", fill="#FFF4D6")
    draw_node(draw, "process", adjust, "edad ← edad − 1", fill="#EDF7E9")
    draw_node(draw, "input", output, "Mostrar edad en años cumplidos", fill="#EAF3F7")
    draw_node(draw, "decision", again, "¿Procesar otro caso?", fill="#FFF4D6")
    draw_node(draw, "terminator", end, "FIN", fill="#EAF3F7")

    draw_arrow(draw, [(850, 305), (850, 375)])
    draw_arrow(draw, [(850, 555), (850, 650)])
    draw_arrow(draw, [(850, 900), (850, 995)], label="Sí", label_at=(875, 930))
    draw_arrow(draw, [(1200, 775), (1250, 775)], label="No", label_at=(1190, 715), color="#A84232")
    draw_arrow(draw, [(1445, 680), (1445, 465), (1350, 465)], color="#A84232")
    draw_arrow(draw, [(850, 1135), (850, 1230)])
    draw_arrow(draw, [(850, 1480), (850, 1580)], label="No", label_at=(875, 1500))
    draw_arrow(draw, [(1190, 1355), (1250, 1355)], label="Sí", label_at=(1180, 1295))
    draw_arrow(draw, [(1440, 1435), (1440, 1645), (1200, 1645)])
    draw_arrow(draw, [(850, 1715), (850, 1810)])
    draw_arrow(draw, [(850, 2030), (850, 2160)], label="No", label_at=(875, 2070))
    draw_arrow(draw, [(560, 1920), (150, 1920), (150, 465), (350, 465)], label="Sí", label_at=(190, 1850))

    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format="PNG", dpi=(180, 180), optimize=True)


def build_leap_diagram(path: Path) -> None:
    image, draw = diagram_base(
        "ALGORITMO 2 · AÑO BISIESTO",
        "Regla gregoriana, validación y repetición controlada",
    )
    start = (610, 210, 1090, 305)
    input_year = (430, 375, 1270, 530)
    positive = (520, 620, 1180, 850)
    error = (1250, 640, 1640, 830)
    residues = (420, 940, 1280, 1095)
    rule = (430, 1180, 1270, 1480)
    yes_output = (120, 1570, 700, 1725)
    no_output = (1000, 1570, 1580, 1725)
    again = (560, 1830, 1140, 2050)
    end = (650, 2160, 1050, 2250)

    draw_node(draw, "terminator", start, "INICIO", fill="#EAF3F7")
    draw_node(draw, "input", input_year, "Ingresar año", fill="#EAF3F7")
    draw_node(draw, "decision", positive, "¿El año es un entero positivo?", fill="#FFF4D6")
    draw_node(draw, "input", error, "Mostrar error", fill="#FCECE8", outline="#A84232")
    draw_node(draw, "process", residues, "Calcular residuos MOD 4, MOD 100 y MOD 400", fill="#EDF7E9")
    draw_node(
        draw,
        "decision",
        rule,
        "¿MOD 400 = 0\no (MOD 4 = 0 y MOD 100 ≠ 0)?",
        fill="#FFF4D6",
    )
    draw_node(draw, "input", yes_output, "Mostrar: sí es bisiesto", fill="#EAF3F7")
    draw_node(draw, "input", no_output, "Mostrar: no es bisiesto", fill="#EAF3F7")
    draw_node(draw, "decision", again, "¿Evaluar otro año?", fill="#FFF4D6")
    draw_node(draw, "terminator", end, "FIN", fill="#EAF3F7")

    draw_arrow(draw, [(850, 305), (850, 375)])
    draw_arrow(draw, [(850, 530), (850, 620)])
    draw_arrow(draw, [(850, 850), (850, 940)], label="Sí", label_at=(875, 875))
    draw_arrow(draw, [(1180, 735), (1250, 735)], label="No", label_at=(1170, 675), color="#A84232")
    draw_arrow(draw, [(1445, 640), (1445, 450), (1270, 450)], color="#A84232")
    draw_arrow(draw, [(850, 1095), (850, 1180)])
    draw_arrow(draw, [(430, 1330), (410, 1330), (410, 1570)], label="Sí", label_at=(300, 1375))
    draw_arrow(draw, [(1270, 1330), (1290, 1330), (1290, 1570)], label="No", label_at=(1310, 1375))
    draw_arrow(draw, [(410, 1725), (410, 1770), (850, 1770), (850, 1830)])
    draw_arrow(draw, [(1290, 1725), (1290, 1770), (850, 1770)])
    draw_arrow(draw, [(850, 2050), (850, 2160)], label="No", label_at=(875, 2080))
    draw_arrow(draw, [(560, 1940), (150, 1940), (150, 450), (430, 450)], label="Sí", label_at=(190, 1870))

    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format="PNG", dpi=(180, 180), optimize=True)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    properties.append(marker)


def set_keep_with_next(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    properties.append(OxmlElement("w:keepNext"))


def set_keep_together(paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    properties.append(OxmlElement("w:keepLines"))


def set_run_font(run, *, size=10.2, bold=False, color=INK, font="Arial", italic=False) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, *, after=4, before=0, line=1.12, alignment=None) -> None:
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment
    set_keep_together(paragraph)


def add_run(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run_font(run, **kwargs)
    return run


def add_page_heading(document, kicker: str, title: str, subtitle: str | None = None) -> None:
    kicker_paragraph = document.add_paragraph()
    format_paragraph(kicker_paragraph, after=1)
    set_keep_with_next(kicker_paragraph)
    add_run(kicker_paragraph, kicker.upper(), size=8.5, bold=True, color=GREEN)

    title_paragraph = document.add_paragraph()
    format_paragraph(title_paragraph, after=3, line=1.0)
    set_keep_with_next(title_paragraph)
    add_run(title_paragraph, title, size=19, bold=True, color=DARK_GREEN)

    if subtitle:
        subtitle_paragraph = document.add_paragraph()
        format_paragraph(subtitle_paragraph, after=8, line=1.05)
        add_run(subtitle_paragraph, subtitle, size=9.4, color=MUTED, italic=True)


def add_subheading(document, text: str) -> None:
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, before=5, after=3)
    set_keep_with_next(paragraph)
    add_run(paragraph, text, size=11.5, bold=True, color=DARK_GREEN)


def add_body(document, text: str, *, bold_lead: str | None = None, after=5) -> None:
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, after=after, line=1.16)
    if bold_lead and text.startswith(bold_lead):
        add_run(paragraph, bold_lead, bold=True)
        add_run(paragraph, text[len(bold_lead):])
    else:
        add_run(paragraph, text)


def add_bullet(document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    format_paragraph(paragraph, after=2, line=1.08)
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.first_line_indent = Inches(-0.15)
    add_run(paragraph, text, size=9.8)


def add_numbered_steps(document, steps: tuple[str, ...]) -> None:
    for number, text in enumerate(steps, start=1):
        paragraph = document.add_paragraph()
        format_paragraph(paragraph, after=2, line=1.08)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        add_run(paragraph, f"{number}. ", size=9.6, bold=True, color=GREEN)
        add_run(paragraph, text, size=9.6)


def add_callout(document, title: str, text: str, *, fill=PALE_GREEN, accent=GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=115, start=150, bottom=115, end=150)
    set_table_borders(table, color=accent, size=10)
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph, after=0, line=1.08)
    add_run(paragraph, f"{title}: ", size=9.5, bold=True, color=DARK_GREEN)
    add_run(paragraph, text, size=9.5)


def add_table(document, headers, rows, *, widths=None, font_size=8.8):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        if widths:
            cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, DARK_GREEN)
        set_cell_margins(cell, top=80, bottom=80)
        paragraph = cell.paragraphs[0]
        format_paragraph(paragraph, after=0, line=1.0)
        add_run(paragraph, str(label), size=font_size, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            if widths:
                cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else CODE_BG)
            set_cell_margins(cell, top=75, bottom=75)
            paragraph = cell.paragraphs[0]
            format_paragraph(paragraph, after=0, line=1.05)
            add_run(paragraph, str(value), size=font_size)
    return table


def add_code_block(document, code: str, *, size=6.75) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=LINE, size=6)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    set_cell_margins(cell, top=95, start=130, bottom=95, end=130)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    lines = code.rstrip().splitlines()
    for index, line in enumerate(lines):
        add_run(paragraph, line or " ", size=size, font="Courier New")
        if index < len(lines) - 1:
            paragraph.add_run().add_break()


def add_picture(document, path: Path, *, width: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(paragraph, after=2, line=1.0)
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), DARK_GREEN)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    properties.append(size)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PÁGINA  ")
    set_run_font(run, size=8, bold=True, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "2"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, value, end):
        run._r.append(node)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(4)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(7.14))
    table.autofit = False
    table.columns[0].width = Inches(3.3)
    table.columns[1].width = Inches(3.84)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, DARK_GREEN)
        set_cell_margins(cell, top=45, start=90, bottom=45, end=90)
    left = table.cell(0, 0).paragraphs[0]
    format_paragraph(left, after=0, line=1.0)
    add_run(left, "SENA · EVIDENCIA TÉCNICA", size=8, bold=True, color=WHITE)
    right = table.cell(0, 1).paragraphs[0]
    format_paragraph(right, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(right, "GA3-220501093-AA2-EV01", size=8, bold=True, color=WHITE)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    add_page_number(footer_paragraph)

    first_footer = section.first_page_footer
    first_paragraph = first_footer.paragraphs[0]
    first_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(first_paragraph, "JHON STEVEN ALVAREZ RUIZ  ·  ANÁLISIS Y DESARROLLO DE SOFTWARE", size=8, bold=True, color=MUTED)


def page_break(document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def format_iso(value: tuple[int, int, int]) -> str:
    year, month, day = value
    return f"{year:04d}-{month:02d}-{day:02d}"


def build_document() -> None:
    age_code = AGE_SOURCE.read_text(encoding="utf-8")
    leap_code = LEAP_SOURCE.read_text(encoding="utf-8")

    document = Document()
    configure_document(document)
    properties = document.core_properties
    properties.title = "Fundamentos de programación estructurada y estructuras cíclicas"
    properties.subject = "GA3-220501093-AA2-EV01 — Algoritmos de edad y año bisiesto"
    properties.author = "Jhon Steven Alvarez Ruiz"
    properties.last_modified_by = "Jhon Steven Alvarez Ruiz"
    properties.keywords = "SENA, algoritmos, pseudocódigo, diagramas de flujo, PSeInt"

    # Página 1 — portada.
    for _ in range(3):
        document.add_paragraph()
    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(banner, color=DARK_GREEN, size=0)
    cell = banner.cell(0, 0)
    set_cell_shading(cell, DEEP_GREEN)
    set_cell_margins(cell, top=180, start=190, bottom=180, end=190)
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(paragraph, "GA3 · PLANEACIÓN · ALGORITMOS", size=11, bold=True, color=WHITE)

    title = document.add_paragraph()
    format_paragraph(title, before=18, after=7, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(title, "FUNDAMENTOS DE\nPROGRAMACIÓN ESTRUCTURADA", size=25, bold=True, color=DARK_GREEN)
    subtitle = document.add_paragraph()
    format_paragraph(subtitle, after=18, line=1.05, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(subtitle, "Edad actual · Año bisiesto · Estructuras de control", size=13, bold=True, color=GREEN)

    cover_table = add_table(
        document,
        ("CAMPO", "INFORMACIÓN"),
        (
            ("Evidencia", "GA3-220501093-AA2-EV01"),
            ("Producto", "Fundamentos de programación estructurada y estructuras cíclicas"),
            ("Aprendiz", "Jhon Steven Alvarez Ruiz"),
            ("Programa", "Análisis y Desarrollo de Software"),
            ("Fase", "Planeación"),
        ),
        widths=(1.55, 5.25),
        font_size=9.2,
    )
    cover_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Página 2 — marco de trabajo.
    page_break(document)
    add_page_heading(
        document,
        "01 · Marco de trabajo",
        "De un problema a una solución verificable",
        "Metodología común para los dos algoritmos solicitados.",
    )
    add_body(
        document,
        "Esta evidencia aplica las fases de análisis, diseño e implementación a dos problemas autónomos. "
        "La solución identifica primero qué información entra y qué resultado debe salir; luego expresa la "
        "lógica con estructuras de control y finalmente la representa mediante seudocódigo y diagramas de flujo.",
    )
    add_subheading(document, "Objetivos")
    add_bullet(document, "Determinar una edad en años cumplidos a partir de dos fechas válidas.")
    add_bullet(document, "Clasificar correctamente un año de acuerdo con la regla gregoriana de años bisiestos.")
    add_bullet(document, "Usar secuencias, decisiones y ciclos con una finalidad concreta y comprobable.")

    add_subheading(document, "Tres fases aplicadas")
    add_table(
        document,
        ("FASE", "PREGUNTA", "RESULTADO EN ESTE DOCUMENTO"),
        (
            ("Análisis", "¿Qué entra, qué se procesa y qué sale?", "Tablas IPO, restricciones y variables."),
            ("Diseño", "¿En qué orden se toman las decisiones?", "Reglas, pasos y diagramas de flujo."),
            ("Implementación", "¿Cómo se expresa de forma ejecutable?", "Seudocódigo compatible con PSeInt y pruebas."),
        ),
        widths=(1.15, 2.45, 3.2),
    )

    add_subheading(document, "Simbología de los diagramas")
    add_table(
        document,
        ("SÍMBOLO", "USO"),
        (
            ("Óvalo", "Inicio o fin del algoritmo."),
            ("Paralelogramo", "Entrada de datos o presentación de resultados."),
            ("Rectángulo", "Proceso, cálculo o asignación."),
            ("Rombo", "Decisión con ramas Sí/No."),
            ("Flecha", "Dirección del flujo y retorno de los ciclos."),
        ),
        widths=(1.5, 5.3),
    )
    add_callout(
        document,
        "Supuesto de calendario",
        "Se usa el calendario gregoriano y el formato YYYY-MM-DD en las pruebas. Para una persona nacida el 29 de febrero, "
        "el cumpleaños se considera cumplido el 1 de marzo durante un año no bisiesto.",
        fill=PALE_BLUE,
        accent="4F8397",
    )

    # Página 3 — análisis del algoritmo 1.
    page_break(document)
    add_page_heading(
        document,
        "02 · Algoritmo 1",
        "Calcular la edad actual",
        "Resultado expresado en años completos, no como una resta aproximada.",
    )
    add_callout(
        document,
        "Problema",
        "A partir de una fecha de nacimiento y una fecha actual, determinar la edad actual de una persona en años.",
    )
    add_subheading(document, "2.1 Análisis: entradas, proceso y salida")
    add_table(
        document,
        ("ENTRADAS", "PROCESO", "SALIDAS"),
        (
            (
                "Día, mes y año de nacimiento.\nDía, mes y año actuales.",
                "Validar ambas fechas; comprobar el orden; restar años; descontar uno si todavía no ocurrió el cumpleaños.",
                "Edad en años cumplidos o mensaje de datos inválidos.",
            ),
        ),
        widths=(2.0, 3.25, 1.55),
        font_size=9.0,
    )
    add_subheading(document, "2.2 Variables y restricciones")
    add_table(
        document,
        ("VARIABLE", "TIPO", "REGLA"),
        (
            ("díaNacimiento / díaActual", "Entero", "Día existente dentro del mes correspondiente."),
            ("mesNacimiento / mesActual", "Entero", "Valor entre 1 y 12."),
            ("añoNacimiento / añoActual", "Entero", "Valor positivo."),
            ("datosVálidos", "Lógico", "Ambas fechas existen y nacimiento ≤ fecha actual."),
            ("edad", "Entero", "Años completos; nunca negativo."),
            ("continuar", "Carácter", "Solo S o N para controlar la repetición."),
        ),
        widths=(2.3, 1.0, 3.5),
        font_size=8.5,
    )
    add_subheading(document, "2.3 Secuencia de solución")
    add_numbered_steps(
        document,
        (
            "Solicitar las dos fechas hasta que sean válidas y estén en orden cronológico.",
            "Calcular la diferencia entre el año actual y el año de nacimiento.",
            "Comparar mes y día para saber si el cumpleaños ya ocurrió en el año actual.",
            "Restar uno cuando la fecha actual todavía sea anterior al cumpleaños.",
            "Mostrar la edad y permitir que el usuario procese otro caso.",
        ),
    )
    add_callout(
        document,
        "Regla central",
        "edad = añoActual − añoNacimiento; si (mesActual, díaActual) < (mesNacimiento, díaNacimiento), entonces edad = edad − 1.",
        fill=PALE_GOLD,
        accent="D09A24",
    )

    # Página 4 — diagrama del algoritmo 1.
    page_break(document)
    add_page_heading(
        document,
        "02 · Algoritmo 1 · Diseño",
        "Diagrama de flujo — calcular la edad",
        "Las flechas de retorno hacen visibles los ciclos de validación y repetición.",
    )
    add_picture(document, AGE_DIAGRAM, width=5.95)
    caption = document.add_paragraph()
    format_paragraph(caption, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(caption, "Figura 1. Flujo completo del algoritmo de edad actual.", size=8.5, italic=True, color=MUTED)

    # Página 5 — seudocódigo del algoritmo 1.
    page_break(document)
    add_page_heading(
        document,
        "02 · Algoritmo 1 · Implementación",
        "Seudocódigo compatible con PSeInt",
        "Incluye validación gregoriana, orden cronológico, cálculo exacto y ciclos controlados.",
    )
    compact_age_code = "\n".join(
        line for line in age_code.rstrip().splitlines() if line.strip()
    )
    add_code_block(document, compact_age_code, size=7.3)

    # Página 6 — pruebas del algoritmo 1.
    page_break(document)
    add_page_heading(
        document,
        "02 · Algoritmo 1 · Verificación",
        "Pruebas de escritorio",
        "Los casos recorren las ramas antes, durante y después del cumpleaños, además de errores de entrada.",
    )
    age_rows = tuple(
        (name, format_iso(birth), format_iso(current), expected)
        for name, birth, current, expected in AGE_TESTS
    )
    add_table(
        document,
        ("CASO", "NACIMIENTO", "FECHA ACTUAL", "RESULTADO"),
        age_rows,
        widths=(2.35, 1.55, 1.55, 1.25),
        font_size=8.2,
    )
    add_subheading(document, "Trazabilidad de dos ramas")
    add_table(
        document,
        ("PASO", "CASO PENDIENTE", "CASO CUMPLIDO"),
        (
            ("Diferencia inicial", "2026 − 2000 = 26", "2026 − 2000 = 26"),
            ("Comparación", "08-16 < 08-20 → Sí", "08-16 < 08-15 → No"),
            ("Ajuste", "26 − 1", "Sin ajuste"),
            ("Salida", "25 años", "26 años"),
        ),
        widths=(1.35, 2.7, 2.7),
        font_size=8.8,
    )
    add_callout(
        document,
        "Resultado de verificación",
        "Las siete pruebas fueron ejecutadas automáticamente por el generador. Las fechas inexistentes o futuras no producen una edad: activan el ciclo de reingreso.",
        fill=PALE_BLUE,
        accent="4F8397",
    )

    # Página 7 — análisis del algoritmo 2.
    page_break(document)
    add_page_heading(
        document,
        "03 · Algoritmo 2",
        "Determinar si un año es bisiesto",
        "La divisibilidad entre 100 y 400 evita clasificar incorrectamente los años seculares.",
    )
    add_callout(
        document,
        "Problema",
        "Determinar si un año indicado es o no un año bisiesto.",
    )
    add_subheading(document, "3.1 Análisis: entradas, proceso y salida")
    add_table(
        document,
        ("ENTRADA", "PROCESO", "SALIDA"),
        (
            (
                "Año entero positivo.",
                "Validar el año; calcular residuos módulo 4, 100 y 400; aplicar la regla gregoriana.",
                "Mensaje: es bisiesto, no es bisiesto o dato inválido.",
            ),
        ),
        widths=(1.6, 3.45, 1.75),
        font_size=9.0,
    )
    add_subheading(document, "3.2 Variables")
    add_table(
        document,
        ("VARIABLE", "TIPO", "FUNCIÓN"),
        (
            ("año", "Entero", "Dato que se desea clasificar; debe ser mayor que cero."),
            ("residuo4 / 100 / 400", "Entero", "Restos usados para evaluar divisibilidad."),
            ("bisiesto", "Lógico", "Verdadero cuando se cumple la regla gregoriana."),
            ("continuar", "Carácter", "Solo S o N para controlar la repetición."),
        ),
        widths=(2.0, 1.15, 3.65),
        font_size=8.7,
    )
    add_subheading(document, "3.3 Tabla de decisión")
    add_table(
        document,
        ("CONDICIÓN", "CLASIFICACIÓN", "EJEMPLO"),
        (
            ("Divisible por 400", "Bisiesto", "2000"),
            ("Divisible por 100, no por 400", "No bisiesto", "1900"),
            ("Divisible por 4, no por 100", "Bisiesto", "2024"),
            ("No divisible por 4", "No bisiesto", "2023"),
        ),
        widths=(3.25, 2.0, 1.55),
        font_size=8.8,
    )
    add_callout(
        document,
        "Expresión lógica",
        "bisiesto = (año MOD 400 = 0) O ((año MOD 4 = 0) Y (año MOD 100 ≠ 0)).",
        fill=PALE_GOLD,
        accent="D09A24",
    )

    # Página 8 — diagrama del algoritmo 2.
    page_break(document)
    add_page_heading(
        document,
        "03 · Algoritmo 2 · Diseño",
        "Diagrama de flujo — año bisiesto",
        "La regla se evalúa después de validar el dato y calcular los tres residuos.",
    )
    add_picture(document, LEAP_DIAGRAM, width=5.95)
    caption = document.add_paragraph()
    format_paragraph(caption, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(caption, "Figura 2. Flujo completo del algoritmo de año bisiesto.", size=8.5, italic=True, color=MUTED)

    # Página 9 — implementación y pruebas del algoritmo 2.
    page_break(document)
    add_page_heading(
        document,
        "03 · Algoritmo 2 · Implementación",
        "Seudocódigo y pruebas",
        "La prueba diferencia años ordinarios, múltiplos de cuatro y excepciones seculares.",
    )
    add_code_block(document, leap_code, size=6.9)
    add_subheading(document, "Pruebas de escritorio")
    add_table(
        document,
        ("AÑO", "RESULTADO", "JUSTIFICACIÓN"),
        LEAP_TESTS,
        widths=(1.0, 1.25, 4.55),
        font_size=8.5,
    )
    add_callout(
        document,
        "Resultado de verificación",
        "Los cinco casos fueron ejecutados automáticamente. El año cero se rechaza y vuelve a solicitarse mediante el ciclo de validación.",
        fill=PALE_BLUE,
        accent="4F8397",
    )

    # Página 10 — síntesis, rúbrica y fuentes.
    page_break(document)
    add_page_heading(
        document,
        "04 · Síntesis",
        "Estructuras usadas y cumplimiento",
        "Correspondencia directa entre la lista de chequeo y la evidencia preparada.",
    )
    add_subheading(document, "4.1 Estructuras de control")
    add_table(
        document,
        ("ESTRUCTURA", "ALGORITMO DE EDAD", "ALGORITMO BISIESTO"),
        (
            ("Secuencial", "Lectura, diferencia de años y salida.", "Cálculo de residuos y salida."),
            ("Condicional", "Validez, orden y cumpleaños pendiente.", "Año positivo y regla de divisibilidad."),
            ("Cíclica", "Reingreso de fechas y procesamiento de otro caso.", "Reingreso del año y evaluación de otro caso."),
        ),
        widths=(1.35, 2.75, 2.7),
        font_size=8.7,
    )
    add_subheading(document, "4.2 Correspondencia con la lista de chequeo")
    add_table(
        document,
        ("INDICADOR", "EVIDENCIA PREPARADA"),
        (
            ("Análisis, diseño e implementación resuelven los problemas.", "Dos soluciones completas, coherentes y verificadas con casos normales y límite."),
            ("Entradas, salidas y procesos para cada algoritmo.", "Tablas IPO, variables, restricciones y secuencias en las páginas 3 y 7."),
            ("Solución mediante diagramas de flujo.", "Dos diagramas con simbología, ramas Sí/No y ciclos visibles en las páginas 4 y 8."),
        ),
        widths=(3.35, 3.45),
        font_size=8.5,
    )
    add_subheading(document, "Conclusiones")
    add_bullet(document, "Una edad exacta requiere comparar mes y día; restar únicamente los años puede producir un resultado adelantado.")
    add_bullet(document, "La regla completa de año bisiesto debe considerar las excepciones de los múltiplos de 100 y 400.")
    add_bullet(document, "Los ciclos de validación evitan continuar con datos imposibles y garantizan una salida coherente.")

    add_subheading(document, "Fuentes consultadas")
    references = (
        ("Guía de aprendizaje GA3 — misma evidencia", "https://archivos.territorio.la/archivos/clases/Guianaprendizajen3nAct___7264160590b5042___.pdf"),
        ("Material formativo oficial SENA — Análisis de algoritmos", "https://zajuna.sena.edu.co/Repositorio/Titulada/institution/SENA/Tecnologia/228118/Contenido/OVA/CF13/index.html"),
        ("National Research Council Canada — Leap years", "https://nrc.canada.ca/en/certifications-evaluations-standards/canadas-official-time/what-years-are-leap-years"),
        ("PSeInt — documentación oficial de pseudocódigo", "https://pseint.sourceforge.net/pseudocodigo.php"),
        ("ISO 8601 — formato de fechas", "https://www.iso.org/iso-8601-date-and-time-format.html"),
    )
    for label, url in references:
        paragraph = document.add_paragraph()
        format_paragraph(paragraph, after=1, line=1.0)
        paragraph.paragraph_format.left_indent = Inches(0.18)
        add_run(paragraph, "• ", size=9, bold=True, color=GREEN)
        add_hyperlink(paragraph, label, url)

    DELIVERY_DIR.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)


def main() -> None:
    run_tests()
    build_age_diagram(AGE_DIAGRAM)
    build_leap_diagram(LEAP_DIAGRAM)
    build_document()
    print(f"Algorithm tests: {len(AGE_TESTS) + len(LEAP_TESTS)} passed")
    print(f"Created: {AGE_DIAGRAM}")
    print(f"Created: {LEAP_DIAGRAM}")
    print(f"Created: {DOCX_PATH}")


if __name__ == "__main__":
    main()
