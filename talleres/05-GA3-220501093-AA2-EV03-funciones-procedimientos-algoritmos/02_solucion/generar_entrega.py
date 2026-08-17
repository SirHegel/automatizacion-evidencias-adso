#!/usr/bin/env python3
"""Genera la entrega pública y, opcionalmente, una variante privada fuera del repositorio."""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import datetime
import json
import math
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
import zipfile

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOLUTION_DIR = Path(__file__).resolve().parent
WORKSHOP_DIR = SOLUTION_DIR.parent
REPO_ROOT = SOLUTION_DIR.parents[2]
PSEUDOCODE_DIR = SOLUTION_DIR / "pseudocodigo"
DIAGRAM_DIR = SOLUTION_DIR / "recursos" / "diagramas"
DELIVERY_DIR = WORKSHOP_DIR / "03_entrega"

EVIDENCE_CODE = "GA3-220501093-AA2-EV03"
PUBLIC_DOCX = DELIVERY_DIR / f"{EVIDENCE_CODE}_Taller_Funciones_Procedimientos.docx"
PUBLIC_ZIP = DELIVERY_DIR / f"{EVIDENCE_CODE}_Taller_Funciones_Procedimientos.zip"

GREEN = "39A900"
DARK_GREEN = "174C2C"
DEEP_GREEN = "0E3322"
PALE_GREEN = "EDF7E9"
PALE_BLUE = "EAF3F7"
PALE_GOLD = "FFF4D6"
PALE_RED = "FCECE8"
INK = "18221B"
MUTED = "5E6A62"
LINE = "D8DED9"
WHITE = "FFFFFF"
CODE_BG = "F4F6F4"

FONT_REGULAR_CANDIDATES = (
    Path("/usr/share/fonts/truetype/msttcorefonts/Arial.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
)
FONT_BOLD_CANDIDATES = (
    Path("/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
)


@dataclass(frozen=True)
class Exercise:
    number: int
    slug: str
    title: str
    statement: str
    inputs: str
    process: str
    output: str
    module: str
    assumptions: tuple[str, ...]
    variables: tuple[tuple[str, str, str], ...]
    tests: tuple[tuple[str, str, str], ...]
    diagram_type: str
    diagram_input: str
    diagram_validation: str
    diagram_process: str
    diagram_output: str
    loop_condition: str = ""

    @property
    def source_path(self) -> Path:
        return PSEUDOCODE_DIR / f"{self.number:02d}_{self.slug}.psc"

    @property
    def diagram_path(self) -> Path:
        return DIAGRAM_DIR / f"{self.number:02d}_{self.slug}.png"

    @property
    def module_name(self) -> str:
        match = re.match(r"^(?:Función|Procedimiento)\s+([A-Za-z][A-Za-z0-9_]*)", self.module)
        if not match:
            raise ValueError(f"No se pudo identificar el módulo documentado: {self.module}")
        return match.group(1)


# ---------------------------------------------------------------------------
# Implementaciones de referencia y pruebas automatizadas
# ---------------------------------------------------------------------------


def marathon_pace(hours: float, minutes: float, distance_km: float) -> float:
    if (
        hours < 0
        or minutes < 0
        or minutes >= 60
        or hours + minutes == 0
        or distance_km <= 0
    ):
        raise ValueError("Tiempo no negativo y distancia positiva requeridos.")
    return (hours * 60 + minutes) / distance_km


def celsius_to_fahrenheit(celsius: float) -> float:
    return (9.0 / 5.0) * celsius + 32


def first_partial_grade(
    workshop_1: float,
    workshop_2: float,
    quiz: float,
    exam: float,
) -> float:
    grades = (workshop_1, workshop_2, quiz, exam)
    if any(grade < 0 or grade > 5 for grade in grades):
        raise ValueError("Cada nota debe estar entre 0 y 5.")
    follow_up = (workshop_1 + workshop_2 + quiz) / 3
    return follow_up * 0.30 + exam * 0.70


def years_to_double(capital: float, annual_rate_percent: float) -> int:
    if capital <= 0 or annual_rate_percent <= 0:
        raise ValueError("El capital y la tasa deben ser positivos.")
    target = 2 * capital
    current = capital
    years = 0
    while current < target:
        current *= 1 + annual_rate_percent / 100
        years += 1
    return years


def numbers_at_most_25(values: list[float]) -> list[float]:
    if len(values) != 20:
        raise ValueError("Se requieren exactamente veinte números.")
    return [value for value in values if value <= 25]


def shirts_total_cop(prices_usd: list[float], exchange_rate: float) -> tuple[float, float]:
    if len(prices_usd) != 5:
        raise ValueError("Se requieren exactamente cinco precios.")
    if any(price < 0 for price in prices_usd) or exchange_rate <= 0:
        raise ValueError("Los precios no pueden ser negativos y la tasa debe ser positiva.")
    total_usd = sum(prices_usd)
    return total_usd, total_usd * exchange_rate


def restaurant_payments(consumptions: list[float]) -> tuple[list[float], float]:
    if not consumptions or any(consumption < 0 for consumption in consumptions):
        raise ValueError("Se requiere al menos un consumo no negativo.")
    payments = [
        consumption * 0.80 if consumption > 50000 else consumption
        for consumption in consumptions
    ]
    return payments, sum(payments)


def next_second(hour: int, minute: int, second: int) -> tuple[int, int, int]:
    if not (0 <= hour <= 23 and 0 <= minute <= 59 and 0 <= second <= 59):
        raise ValueError("La hora ingresada no es válida.")
    second += 1
    if second == 60:
        second = 0
        minute += 1
    if minute == 60:
        minute = 0
        hour += 1
    if hour == 24:
        hour = 0
    return hour, minute, second


def product_one_to_n(number: int) -> int:
    if number < 0:
        raise ValueError("N debe ser un entero no negativo.")
    product = 1
    for factor in range(1, number + 1):
        product *= factor
    return product


def descending_multiplication_table(number: int) -> list[tuple[int, int]]:
    if number < 1 or number > 10:
        raise ValueError("El número debe estar entre 1 y 10.")
    return [(factor, number * factor) for factor in range(10, 0, -1)]


def run_tests() -> int:
    checks = 0

    def equal(actual, expected) -> None:
        nonlocal checks
        if actual != expected:
            raise AssertionError(f"Resultado inesperado: {actual!r} != {expected!r}")
        checks += 1

    def close(actual: float, expected: float, tolerance: float = 1e-8) -> None:
        nonlocal checks
        if not math.isclose(actual, expected, rel_tol=tolerance, abs_tol=tolerance):
            raise AssertionError(f"Resultado inesperado: {actual!r} != {expected!r}")
        checks += 1

    def rejects(function, *arguments) -> None:
        nonlocal checks
        try:
            function(*arguments)
        except ValueError:
            checks += 1
            return
        raise AssertionError(f"{function.__name__} debía rechazar {arguments!r}")

    close(marathon_pace(2, 25, 42.195), 145 / 42.195)
    close(marathon_pace(0, 30, 10), 3)
    rejects(marathon_pace, 1, 0, 0)
    rejects(marathon_pace, 0, 0, 42.195)
    rejects(marathon_pace, 1, 60, 42.195)

    close(celsius_to_fahrenheit(0), 32)
    close(celsius_to_fahrenheit(100), 212)
    close(celsius_to_fahrenheit(-40), -40)

    close(first_partial_grade(4, 3.5, 4.5, 3.8), 3.86)
    close(first_partial_grade(5, 5, 5, 5), 5)
    rejects(first_partial_grade, 6, 4, 4, 4)

    equal(years_to_double(100000, 10), 8)
    equal(years_to_double(100, 100), 1)
    rejects(years_to_double, 100, 0)

    sample_numbers = list(range(1, 21))
    equal(numbers_at_most_25(sample_numbers), sample_numbers)
    equal(numbers_at_most_25(list(range(20, 40))), list(range(20, 26)))
    rejects(numbers_at_most_25, [1, 2])

    equal(shirts_total_cop([20, 25, 30, 15, 10], 4000), (100, 400000))
    equal(shirts_total_cop([0, 0, 0, 0, 0], 4100), (0, 0))
    rejects(shirts_total_cop, [10, 20], 4000)

    equal(restaurant_payments([50000, 60000]), ([50000, 48000], 98000))
    equal(restaurant_payments([0]), ([0], 0))
    rejects(restaurant_payments, [-1])

    equal(next_second(12, 30, 45), (12, 30, 46))
    equal(next_second(12, 59, 59), (13, 0, 0))
    equal(next_second(23, 59, 59), (0, 0, 0))
    rejects(next_second, 24, 0, 0)

    equal(product_one_to_n(5), 120)
    equal(product_one_to_n(0), 1)
    rejects(product_one_to_n, -1)

    table = descending_multiplication_table(3)
    equal(table[0], (10, 30))
    equal(table[-1], (1, 3))
    equal(len(table), 10)
    rejects(descending_multiplication_table, 11)
    return checks


EXERCISES = (
    Exercise(
        1,
        "ritmo_maraton",
        "Tiempo medio de una maratón",
        "Un corredor completa 42,195 km en 2 horas y 25 minutos. Calcular el tiempo medio en minutos por kilómetro.",
        "Horas, minutos y distancia en kilómetros.",
        "Convertir todo el tiempo a minutos y dividirlo entre la distancia.",
        "Ritmo medio en minutos por kilómetro.",
        "Función CalcularRitmoMinutosPorKilometro(distancia, horas, minutos)",
        (
            "La distancia debe ser positiva y los tiempos no pueden ser negativos.",
            "La función se parametriza aunque el caso de la guía proporcione valores fijos.",
        ),
        (
            ("horas / minutos", "Entero", "Duración de la carrera."),
            ("distancia", "Real", "Recorrido en kilómetros; mayor que cero."),
            ("ritmo", "Real", "Minutos empleados por kilómetro."),
        ),
        (
            ("Caso guía", "2 h, 25 min, 42,195 km", "3,436426 min/km"),
            ("Control", "0 h, 30 min, 10 km", "3 min/km"),
            ("Inválido", "Distancia = 0", "Solicitar nuevamente"),
        ),
        "simple",
        "Leer distancia, horas y minutos",
        "¿distancia > 0, horas ≥ 0, 0 ≤ minutos < 60 y tiempo > 0?",
        "ritmo ← CalcularRitmoMinutosPorKilometro(distancia, horas, minutos)",
        "Mostrar ritmo en min/km",
    ),
    Exercise(
        2,
        "celsius_fahrenheit",
        "Conversión de Celsius a Fahrenheit",
        "Convertir una temperatura dada en grados Celsius a grados Fahrenheit mediante F = (9/5)C + 32.",
        "Temperatura real en grados Celsius.",
        "Multiplicar por 9/5 y sumar 32.",
        "Temperatura equivalente en grados Fahrenheit.",
        "Función ConvertirCelsiusAFahrenheit(celsius)",
        ("Se aceptan temperaturas positivas, negativas y cero.",),
        (
            ("celsius", "Real", "Temperatura de origen."),
            ("fahrenheit", "Real", "Resultado de la conversión."),
        ),
        (
            ("Congelación", "0 °C", "32 °F"),
            ("Ebullición", "100 °C", "212 °F"),
            ("Coincidencia", "−40 °C", "−40 °F"),
        ),
        "simple_no_validation",
        "Leer temperatura Celsius",
        "",
        "fahrenheit ← ConvertirCelsiusAFahrenheit(celsius)",
        "Mostrar temperatura Fahrenheit",
    ),
    Exercise(
        3,
        "nota_primer_parcial",
        "Nota del primer parcial",
        "Calcular la nota del primer parcial: dos talleres y un cuestionario representan conjuntamente el 30 %, y el examen el 70 %.",
        "Notas de taller 1, taller 2, cuestionario y examen.",
        "Promediar las tres actividades; ponderar 30 % y sumar el examen ponderado al 70 %.",
        "Nota definitiva del primer parcial.",
        "Función CalcularNotaPrimerParcial(taller1, taller2, quiz, examen)",
        (
            "Se adopta la escala de 0 a 5 para todas las notas.",
            "Las tres actividades tienen el mismo peso dentro del componente del 30 %.",
        ),
        (
            ("taller1 / taller2", "Real", "Notas de los dos talleres."),
            ("quiz / examen", "Real", "Notas del cuestionario y examen."),
            ("promedioSeguimiento", "Real", "Promedio de las tres actividades."),
            ("notaFinal", "Real", "Resultado ponderado."),
        ),
        (
            ("Caso base", "4; 3,5; 4,5; 3,8", "3,86"),
            ("Máximo", "5; 5; 5; 5", "5,00"),
            ("Inválido", "Alguna nota > 5", "Solicitar nuevamente"),
        ),
        "simple",
        "Leer dos talleres, quiz y examen",
        "¿Todas las notas están entre 0 y 5?",
        "notaFinal ← CalcularNotaPrimerParcial(taller1, taller2, quiz, examen)",
        "Mostrar nota definitiva",
    ),
    Exercise(
        4,
        "duplicar_capital",
        "Años necesarios para duplicar un capital",
        "Dado un capital C y una tasa anual R, determinar el primer año en el que el capital alcanza al menos el doble.",
        "Capital inicial y tasa anual expresada como porcentaje.",
        "Aplicar capitalización compuesta cada año hasta alcanzar 2 × C.",
        "Cantidad entera de años necesarios.",
        "Procedimiento CalcularDuplicacion (salidas por referencia)",
        (
            "Se usa interés compuesto con capitalización anual.",
            "Capital y tasa deben ser estrictamente positivos.",
        ),
        (
            ("capital / meta", "Real", "Valor inicial y objetivo equivalente al doble."),
            ("tasaAnual", "Real", "Porcentaje anual positivo."),
            ("montoFinal", "Real", "Capital acumulado durante el ciclo."),
            ("anios", "Entero", "Contador de capitalizaciones."),
        ),
        (
            ("Caso base", "C = 100000; R = 10 %", "8 años"),
            ("Tasa 100 %", "C = 100; R = 100 %", "1 año"),
            ("Inválido", "R = 0", "Solicitar nuevamente"),
        ),
        "loop",
        "Leer capital y tasaAnual",
        "¿capital > 0 y tasaAnual > 0?",
        "montoFinal ← montoFinal × (1 + tasaAnual/100); anios ← anios + 1",
        "Mostrar años",
        "¿montoFinal < 2 × capital?",
    ),
    Exercise(
        5,
        "numeros_menores_igual_25",
        "Filtrado de veinte números",
        "Ingresar veinte números y mostrar todos los valores menores o iguales a 25.",
        "Veinte números reales.",
        "Repetir veinte lecturas; evaluar cada número con una función y mostrarlo si cumple.",
        "Valores menores o iguales a 25 y cantidad encontrada.",
        "Función EsMenorOIgualAlLimite(numero, limite)",
        ("Se procesan exactamente veinte entradas, incluyendo valores repetidos.",),
        (
            ("numero / limite", "Real", "Valor leído y frontera fija igual a 25."),
            ("posicion / totalDatos", "Entero", "Controlan las veinte lecturas."),
            ("cantidadMostrada", "Entero", "Cuenta los valores que cumplen."),
        ),
        (
            ("Todos cumplen", "Valores 1 a 20", "Se muestran 20"),
            ("Caso mixto", "Valores 20 a 39", "Se muestran 20 a 25"),
            ("Frontera", "Valor = 25", "Sí se muestra"),
        ),
        "loop_no_validation",
        "Preparar veinte lecturas",
        "",
        "Leer número; EvaluarYMostrar(numero, limite, cantidadMostrada); posicion ← posicion + 1",
        "Mostrar cantidad encontrada",
        "¿posicion ≤ totalDatos?",
    ),
    Exercise(
        6,
        "camisas_dolares_pesos",
        "Venta de cinco camisas",
        "Sumar cinco precios de camisas en dólares y mostrar el total de la venta convertido a pesos.",
        "Cinco precios en dólares y tasa de cambio en pesos por dólar.",
        "Validar la tasa y los precios, sumar cinco valores y convertir el total.",
        "Total en dólares y total equivalente en pesos.",
        "Función ConvertirDolaresAPesos(totalDolares, tasaCambio)",
        (
            "La tasa de cambio es una entrada positiva para evitar fijar un valor que cambia con el tiempo.",
            "Los precios pueden ser cero, pero no negativos.",
        ),
        (
            ("precio", "Real", "Precio de cada camisa en dólares."),
            ("tasaCambio", "Real", "Pesos equivalentes a un dólar."),
            ("totalDolares / totalPesos", "Real", "Acumulados antes y después de convertir."),
            ("posicion", "Entero", "Controla cinco lecturas."),
        ),
        (
            ("Caso base", "20, 25, 30, 15, 10; tasa 4000", "USD 100; COP 400000"),
            ("Precios cero", "Cinco ceros; tasa 4100", "USD 0; COP 0"),
            ("Inválido", "Tasa ≤ 0", "Solicitar nuevamente"),
        ),
        "loop",
        "Leer tasaCambio",
        "¿tasaCambio > 0?",
        "LeerPrecioValido(posicion, precio); acumular totalDolares; posicion ← posicion + 1",
        "Convertir total y mostrar USD/COP",
        "¿posicion ≤ 5?",
    ),
    Exercise(
        7,
        "consumos_restaurante",
        "Pagos de clientes de un restaurante",
        "Registrar consumos; aplicar 20 % de descuento cuando cada consumo exceda 50000; mostrar cada pago y el total.",
        "Cantidad de clientes y consumo no negativo de cada uno.",
        "Calcular el pago individual, mostrarlo y acumular todos los pagos.",
        "Pago de cada cliente y total general.",
        "Función CalcularPago(consumo)",
        (
            "La cantidad de clientes se solicita antes de iniciar y debe ser positiva.",
            "Un consumo exactamente igual a 50000 no recibe descuento.",
        ),
        (
            ("cantidadClientes", "Entero", "Número de clientes por procesar."),
            ("consumo / pago", "Real", "Valor original y valor después del descuento."),
            ("totalPagos", "Real", "Acumulador de pagos individuales."),
            ("cliente", "Entero", "Control del recorrido de clientes."),
        ),
        (
            ("Frontera", "Consumo 50000", "Pago 50000"),
            ("Con descuento", "Consumo 60000", "Pago 48000"),
            ("Total", "Consumos 50000 y 60000", "Total 98000"),
        ),
        "loop",
        "Leer cantidadClientes",
        "¿cantidadClientes > 0?",
        "LeerConsumoValido(cliente, consumo); CalcularPago; mostrar y acumular; cliente ← cliente + 1",
        "Mostrar total de pagos",
        "¿cliente ≤ cantidadClientes?",
    ),
    Exercise(
        8,
        "siguiente_segundo",
        "Hora del siguiente segundo",
        "Ingresar hora, minutos y segundos válidos y calcular la hora correspondiente al siguiente segundo.",
        "Hora de 0 a 23; minutos y segundos de 0 a 59.",
        "Incrementar segundos y propagar los cambios a minutos y horas, incluyendo el cambio de día.",
        "Hora, minuto y segundo posteriores.",
        "Procedimiento AvanzarUnSegundo (parámetros por referencia)",
        ("Después de 23:59:59 se obtiene 00:00:00.",),
        (
            ("hora", "Entero", "Componente horaria; rango 0..23."),
            ("minuto / segundo", "Entero", "Componentes; rango 0..59."),
            ("valida", "Lógico", "Controla la aceptación de la entrada."),
        ),
        (
            ("Normal", "12:30:45", "12:30:46"),
            ("Cambio de hora", "12:59:59", "13:00:00"),
            ("Cambio de día", "23:59:59", "00:00:00"),
        ),
        "simple",
        "Leer hora, minuto y segundo",
        "¿0≤hora≤23, 0≤minuto≤59 y 0≤segundo≤59?",
        "AvanzarUnSegundo(hora, minuto, segundo): incrementar y propagar acarreos",
        "Mostrar HH:MM:SS",
    ),
    Exercise(
        9,
        "producto_1_hasta_n",
        "Producto desde 1 hasta N",
        "Dado un número entero N, calcular el producto de todos los enteros desde 1 hasta N.",
        "Número entero N mayor o igual a cero.",
        "Inicializar el producto en 1 y multiplicar sucesivamente desde 1 hasta N.",
        "Producto acumulado, equivalente a N factorial.",
        "Función CalcularProductoUnoHastaN(n)",
        (
            "Se acepta N = 0 y se aplica la convención 0! = 1.",
            "Para evitar desbordamiento se recomienda trabajar con valores pequeños.",
        ),
        (
            ("numeroIngresado / producto", "Real", "Dato leído y acumulador inicializado en 1."),
            ("n / factor", "Entero", "Límite validado y contador desde 1."),
        ),
        (
            ("Caso base", "N = 5", "120"),
            ("Caso cero", "N = 0", "1"),
            ("Inválido", "N < 0", "Solicitar nuevamente"),
        ),
        "loop",
        "Leer numeroIngresado",
        "¿numeroIngresado es entero y ≥ 0?",
        "producto ← producto × factor; factor ← factor + 1",
        "Mostrar producto",
        "¿factor ≤ N?",
    ),
    Exercise(
        10,
        "tabla_multiplicar_decreciente",
        "Tabla de multiplicar decreciente",
        "Mostrar la tabla de multiplicar decreciente de un número ingresado entre 1 y 10.",
        "Número entero en el rango de 1 a 10.",
        "Recorrer los multiplicadores desde 10 hasta 1 y mostrar cada producto.",
        "Diez filas de la tabla en orden descendente.",
        "Procedimiento MostrarTablaDecreciente(numero)",
        ("El multiplicador disminuye de uno en uno: 10, 9, …, 1.",),
        (
            ("numeroIngresado", "Real", "Dato leído antes de validar su integralidad."),
            ("numero", "Entero", "Base de la tabla; rango 1..10."),
            ("multiplicador", "Entero", "Contador descendente desde 10."),
        ),
        (
            ("Primera fila", "Número 3; factor 10", "3 × 10 = 30"),
            ("Última fila", "Número 3; factor 1", "3 × 1 = 3"),
            ("Inválido", "Número 11", "Solicitar nuevamente"),
        ),
        "loop",
        "Leer numeroIngresado",
        "¿numeroIngresado es entero entre 1 y 10?",
        "Mostrar número × multiplicador; multiplicador ← multiplicador − 1",
        "Finalizar tabla",
        "¿multiplicador ≥ 1?",
    ),
)


def validate_traceability() -> None:
    """Evita publicar diagramas que nombren módulos inexistentes en PSeInt."""
    expected_paths = {exercise.source_path.resolve() for exercise in EXERCISES}
    actual_paths = {path.resolve() for path in PSEUDOCODE_DIR.glob("*.psc")}
    if actual_paths != expected_paths:
        missing = sorted(str(path) for path in expected_paths - actual_paths)
        unexpected = sorted(str(path) for path in actual_paths - expected_paths)
        raise ValueError(f"Inventario PSeInt incoherente; faltan={missing}; sobran={unexpected}")

    for exercise in EXERCISES:
        source = exercise.source_path.read_text(encoding="utf-8")
        module_pattern = re.compile(
            rf"^SubProceso(?:\s+\w+\s+<-)?\s+{re.escape(exercise.module_name)}\s*\(",
            flags=re.MULTILINE,
        )
        if not module_pattern.search(source):
            raise ValueError(
                f"Problema {exercise.number}: el módulo documentado "
                f"{exercise.module_name} no existe en {exercise.source_path.name}."
            )


# ---------------------------------------------------------------------------
# Diagramas de flujo
# ---------------------------------------------------------------------------


def first_existing(candidates: tuple[Path, ...]) -> Path:
    for path in candidates:
        if path.is_file():
            return path
    raise FileNotFoundError(f"No se encontró ninguna tipografía: {candidates}")


def image_font(size: int, *, bold: bool = False):
    candidates = FONT_BOLD_CANDIDATES if bold else FONT_REGULAR_CANDIDATES
    return ImageFont.truetype(str(first_existing(candidates)), size=size)


def wrap_for_width(draw, text: str, font, max_width: int) -> str:
    lines: list[str] = []
    for original_line in text.splitlines() or [""]:
        words = original_line.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return "\n".join(lines)


def draw_centered_text(draw, bounds, text: str, *, font, fill="#18221B", padding=28):
    left, top, right, bottom = bounds
    wrapped = wrap_for_width(draw, text, font, right - left - 2 * padding)
    box = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=6, align="center")
    width = box[2] - box[0]
    height = box[3] - box[1]
    x = (left + right - width) / 2
    y = (top + bottom - height) / 2 - box[1]
    draw.multiline_text((x, y), wrapped, font=font, fill=fill, spacing=6, align="center")


def draw_node(draw, kind: str, bounds, text: str, *, fill: str, outline="#174C2C"):
    left, top, right, bottom = bounds
    if kind == "terminator":
        draw.rounded_rectangle(bounds, radius=(bottom - top) // 2, fill=fill, outline=outline, width=5)
    elif kind == "input":
        offset = min(45, (right - left) // 8)
        points = [(left + offset, top), (right, top), (right - offset, bottom), (left, bottom)]
        draw.polygon(points, fill=fill)
        draw.line(points + [points[0]], fill=outline, width=5, joint="curve")
    elif kind == "decision":
        points = [
            ((left + right) // 2, top),
            (right, (top + bottom) // 2),
            ((left + right) // 2, bottom),
            (left, (top + bottom) // 2),
        ]
        draw.polygon(points, fill=fill)
        draw.line(points + [points[0]], fill=outline, width=5, joint="curve")
    else:
        draw.rounded_rectangle(bounds, radius=16, fill=fill, outline=outline, width=5)
    font_size = 30 if len(text) < 80 else 26
    draw_centered_text(draw, bounds, text, font=image_font(font_size, bold=True), padding=48)


def draw_arrow(draw, points, *, label: str | None = None, label_at=None, color="#174C2C"):
    draw.line(points, fill=color, width=6, joint="curve")
    start = points[-2]
    end = points[-1]
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 22
    spread = math.pi / 6
    arrow = [
        end,
        (end[0] - length * math.cos(angle - spread), end[1] - length * math.sin(angle - spread)),
        (end[0] - length * math.cos(angle + spread), end[1] - length * math.sin(angle + spread)),
    ]
    draw.polygon(arrow, fill=color)
    if label and label_at:
        font = image_font(24, bold=True)
        box = draw.textbbox((0, 0), label, font=font)
        width = box[2] - box[0]
        height = box[3] - box[1]
        x, y = label_at
        draw.rounded_rectangle(
            (x - 10, y - 6, x + width + 10, y + height + 6),
            radius=8,
            fill="#FFFFFF",
            outline="#D8DED9",
            width=2,
        )
        draw.text((x, y), label, font=font, fill=color)


def diagram_base(exercise: Exercise, height: int):
    image = Image.new("RGB", (1800, height), "#F7FAF7")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((55, 35, 1745, 125), radius=20, fill="#174C2C")
    draw.text(
        (90, 53),
        f"PROBLEMA {exercise.number:02d} · {exercise.title.upper()}",
        font=image_font(38, bold=True),
        fill="#FFFFFF",
    )
    draw.text(
        (85, 142),
        f"Módulo principal: {exercise.module}",
        font=image_font(24),
        fill="#5E6A62",
    )
    return image, draw


def build_simple_diagram(exercise: Exercise) -> None:
    has_validation = exercise.diagram_type != "simple_no_validation"
    image, draw = diagram_base(exercise, 1180)
    start = (690, 205, 1110, 275)
    input_node = (410, 325, 1390, 430)
    decision = (550, 475, 1250, 650)
    error = (1350, 505, 1740, 625)
    process = (410, 725 if has_validation else 535, 1390, 845 if has_validation else 655)
    output = (430, 900 if has_validation else 735, 1370, 1010 if has_validation else 845)
    end = (710, 1060 if has_validation else 935, 1090, 1130 if has_validation else 1005)

    draw_node(draw, "terminator", start, "INICIO", fill="#EAF3F7")
    draw_node(draw, "input", input_node, exercise.diagram_input, fill="#EAF3F7")
    draw_arrow(draw, [(900, 275), (900, 325)])

    if has_validation:
        draw_node(draw, "decision", decision, exercise.diagram_validation, fill="#FFF4D6")
        draw_node(draw, "input", error, "Mostrar error y solicitar nuevamente", fill="#FCECE8", outline="#A84232")
        draw_arrow(draw, [(900, 430), (900, 475)])
        draw_arrow(draw, [(900, 650), (900, 725)], label="Sí", label_at=(925, 670))
        draw_arrow(draw, [(1250, 562), (1350, 562)], label="No", label_at=(1260, 515), color="#A84232")
        draw_arrow(draw, [(1545, 505), (1545, 378), (1390, 378)], color="#A84232")
    else:
        draw_arrow(draw, [(900, 430), (900, 535)])

    draw_node(draw, "process", process, exercise.diagram_process, fill="#EDF7E9")
    draw_node(draw, "input", output, exercise.diagram_output, fill="#EAF3F7")
    draw_node(draw, "terminator", end, "FIN", fill="#EAF3F7")
    draw_arrow(draw, [(900, process[3]), (900, output[1])])
    draw_arrow(draw, [(900, output[3]), (900, end[1])])
    exercise.diagram_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(exercise.diagram_path, optimize=True)


def build_loop_diagram(exercise: Exercise) -> None:
    has_validation = exercise.diagram_type != "loop_no_validation"
    image, draw = diagram_base(exercise, 1300)
    start = (700, 200, 1100, 270)
    input_node = (410, 315, 1390, 420)
    validation = (555, 465, 1245, 625)
    error = (1360, 490, 1740, 610)
    initialize = (450, 680 if has_validation else 490, 1350, 775 if has_validation else 585)
    condition = (565, 820 if has_validation else 635, 1235, 990 if has_validation else 805)
    body = (1300, 845 if has_validation else 660, 1760, 965 if has_validation else 780)
    output = (470, 1050 if has_validation else 865, 1330, 1145 if has_validation else 960)
    end = (720, 1190 if has_validation else 1010, 1080, 1260 if has_validation else 1080)

    draw_node(draw, "terminator", start, "INICIO", fill="#EAF3F7")
    draw_node(draw, "input", input_node, exercise.diagram_input, fill="#EAF3F7")
    draw_arrow(draw, [(900, 270), (900, 315)])

    if has_validation:
        draw_node(draw, "decision", validation, exercise.diagram_validation, fill="#FFF4D6")
        draw_node(draw, "input", error, "Mostrar error y solicitar nuevamente", fill="#FCECE8", outline="#A84232")
        draw_arrow(draw, [(900, 420), (900, 465)])
        draw_arrow(draw, [(900, 625), (900, 680)], label="Sí", label_at=(925, 640))
        draw_arrow(draw, [(1245, 545), (1360, 545)], label="No", label_at=(1260, 500), color="#A84232")
        draw_arrow(draw, [(1550, 490), (1550, 365), (1390, 365)], color="#A84232")
    else:
        draw_arrow(draw, [(900, 420), (900, 490)])

    draw_node(draw, "process", initialize, "Inicializar contadores y acumuladores", fill="#EDF7E9")
    draw_node(draw, "decision", condition, exercise.loop_condition, fill="#FFF4D6")
    draw_node(draw, "process", body, exercise.diagram_process, fill="#EDF7E9")
    draw_node(draw, "input", output, exercise.diagram_output, fill="#EAF3F7")
    draw_node(draw, "terminator", end, "FIN", fill="#EAF3F7")
    draw_arrow(draw, [(900, initialize[3]), (900, condition[1])])
    draw_arrow(draw, [(1235, (condition[1] + condition[3]) // 2), (1300, (body[1] + body[3]) // 2)], label="Sí", label_at=(1225, condition[1] + 15))
    draw_arrow(draw, [(1530, body[1]), (1530, condition[1] - 30), (900, condition[1] - 30), (900, condition[1])])
    draw_arrow(draw, [(900, condition[3]), (900, output[1])], label="No", label_at=(925, condition[3] + 10))
    draw_arrow(draw, [(900, output[3]), (900, end[1])])
    exercise.diagram_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(exercise.diagram_path, optimize=True)


def build_diagrams() -> None:
    for exercise in EXERCISES:
        if exercise.diagram_type.startswith("simple"):
            build_simple_diagram(exercise)
        else:
            build_loop_diagram(exercise)


# ---------------------------------------------------------------------------
# Documento académico
# ---------------------------------------------------------------------------


def set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, *, top=70, start=90, bottom=70, end=90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, color=LINE, size=5) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def format_paragraph(paragraph, *, before=0, after=4, line=1.08, alignment=None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if alignment is not None:
        paragraph.alignment = alignment


def add_run(
    paragraph,
    text: str,
    *,
    size=9.5,
    bold=False,
    italic=False,
    color=INK,
    font="Arial",
):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def configure_document(document: Document, *, author: str, private: bool) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.18)
    section.footer_distance = Inches(0.20)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.5)

    properties = document.core_properties
    properties.title = "Taller funciones y procedimientos en la solución de algoritmos"
    properties.subject = EVIDENCE_CODE
    properties.author = author
    properties.last_modified_by = author
    properties.keywords = "algoritmos, funciones, procedimientos, pseudocódigo, diagramas de flujo"
    properties.comments = "Documento privado local" if private else "Versión pública anonimizada"
    fixed_date = datetime(2026, 8, 16, 12, 0, 0)
    properties.created = fixed_date
    properties.modified = fixed_date

    header = section.header
    paragraph = header.paragraphs[0]
    format_paragraph(paragraph, after=0, line=1.0)
    table = header.add_table(rows=1, cols=2, width=Inches(7.14))
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(2.34)
    set_table_borders(table, color=DARK_GREEN, size=0)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, DARK_GREEN)
        set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    format_paragraph(left, after=0, line=1.0)
    format_paragraph(right, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(left, "SENA · EVIDENCIA TÉCNICA", size=8.2, bold=True, color=WHITE)
    add_run(right, EVIDENCE_CODE, size=8.2, bold=True, color=WHITE)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    format_paragraph(paragraph, after=0, line=1.0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(paragraph, "PÁGINA ", size=8, bold=True, color=MUTED)
    add_page_field(paragraph)


def page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_page_heading(document, eyebrow: str, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, after=1, line=1.0)
    add_run(paragraph, eyebrow.upper(), size=9, bold=True, color=GREEN)
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, after=2, line=1.0)
    add_run(paragraph, title, size=22, bold=True, color=DARK_GREEN)
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, after=8, line=1.0)
    add_run(paragraph, subtitle, size=10.3, italic=True, color=MUTED)


def add_subheading(document, text: str) -> None:
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, before=3, after=3, line=1.0)
    add_run(paragraph, text, size=12.2, bold=True, color=DARK_GREEN)


def add_bullet(document, text: str) -> None:
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, after=2, line=1.05)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.first_line_indent = Inches(-0.14)
    add_run(paragraph, "• ", size=9.4, bold=True, color=GREEN)
    add_run(paragraph, text, size=9.4)


def add_callout(document, label: str, text: str, *, fill=PALE_GREEN, accent=GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=accent, size=7)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph, after=0, line=1.08)
    add_run(paragraph, f"{label}: ", size=9.5, bold=True, color=DARK_GREEN)
    add_run(paragraph, text, size=9.5)


def add_table(document, headers, rows, *, widths=None, font_size=8.4):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, label in enumerate(headers):
        cell = header.cells[index]
        if widths:
            cell.width = Inches(widths[index])
        set_cell_shading(cell, DARK_GREEN)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        format_paragraph(paragraph, after=0, line=1.0)
        add_run(paragraph, str(label), size=font_size, bold=True, color=WHITE)
    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cell = cells[index]
            if widths:
                cell.width = Inches(widths[index])
            if row_index % 2:
                set_cell_shading(cell, "F3F6F3")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            format_paragraph(paragraph, after=0, line=1.0)
            add_run(paragraph, str(value), size=font_size)
    return table


def add_code_block(document, code: str) -> None:
    compact_lines = [line for line in code.rstrip().splitlines() if line.strip()]
    line_count = len(compact_lines)
    size = 8.4 if line_count <= 40 else 7.35 if line_count <= 54 else 6.8 if line_count <= 64 else 6.25
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=LINE, size=6)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    set_cell_margins(cell, top=80, start=105, bottom=80, end=105)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for index, line in enumerate(compact_lines):
        add_run(paragraph, line, size=size, font="Courier New")
        if index < line_count - 1:
            paragraph.add_run().add_break()


def add_picture(document, path: Path, *, width=6.55) -> None:
    paragraph = document.add_paragraph()
    format_paragraph(paragraph, after=2, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def build_cover(document: Document, *, profile: dict[str, str] | None) -> None:
    for _ in range(3):
        document.add_paragraph()
    label = document.add_paragraph()
    format_paragraph(label, after=14, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(label, "EVIDENCIA DE PRODUCTO", size=11, bold=True, color=GREEN)

    title = document.add_paragraph()
    format_paragraph(title, after=10, line=1.02, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(
        title,
        "Taller funciones y procedimientos\nen la solución de algoritmos",
        size=27,
        bold=True,
        color=DARK_GREEN,
    )
    code = document.add_paragraph()
    format_paragraph(code, after=26, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(code, EVIDENCE_CODE, size=13, bold=True, color=MUTED)

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=GREEN, size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GREEN if profile else PALE_BLUE)
    set_cell_margins(cell, top=180, start=180, bottom=180, end=180)
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph, after=0, line=1.15, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    if profile:
        add_run(paragraph, "INFORME COMPLETO LOCAL · NO PUBLICAR\n", size=10, bold=True, color="A84232")
        add_run(paragraph, profile["nombre_completo"], size=15, bold=True, color=DARK_GREEN)
        for field in ("programa", "institucion", "centro", "ciudad", "fecha"):
            value = str(profile.get(field, "")).strip()
            if value:
                add_run(paragraph, f"\n{value}", size=10, color=MUTED)
    else:
        add_run(paragraph, "VERSIÓN PÚBLICA ANONIMIZADA", size=13, bold=True, color=DARK_GREEN)
        add_run(paragraph, "\nSin identidad, contacto ni datos académicos individualizantes", size=9.5, color=MUTED)

    document.add_paragraph()
    summary = document.add_paragraph()
    format_paragraph(summary, after=0, line=1.18, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(
        summary,
        "10 problemas · 10 pseudocódigos · 10 diagramas de flujo\nFunciones, procedimientos, decisiones y ciclos verificables",
        size=11,
        color=INK,
    )


def build_document(path: Path, *, profile: dict[str, str] | None = None) -> None:
    document = Document()
    author = profile["nombre_completo"] if profile else "Entrega académica pública"
    configure_document(document, author=author, private=bool(profile))
    build_cover(document, profile=profile)

    page_break(document)
    add_page_heading(
        document,
        "00 · Arquitectura de la solución",
        "Cómo leer esta entrega",
        "Cada problema se analiza, representa y comprueba con la misma trazabilidad.",
    )
    add_callout(
        document,
        "Cobertura",
        "La entrega contiene exactamente diez pseudocódigos y diez diagramas. Cada pseudocódigo separa el proceso principal de una función o procedimiento reutilizable.",
    )
    add_subheading(document, "Método aplicado")
    add_table(
        document,
        ("FASE", "CONTENIDO", "RESULTADO VERIFICABLE"),
        (
            ("Análisis", "Enunciado, entradas, proceso, salida, variables y supuestos.", "Tabla IPO y casos de prueba."),
            ("Diseño", "Secuencia, decisiones y ciclos con simbología consistente.", "Un diagrama de flujo por problema."),
            ("Implementación", "Función o procedimiento y proceso principal.", "Un archivo PSeInt por problema."),
            ("Verificación", "Casos normales, frontera y entradas inválidas.", "Pruebas de escritorio y pruebas ejecutables."),
        ),
        widths=(1.2, 3.2, 2.6),
        font_size=8.8,
    )
    add_subheading(document, "Simbología")
    add_table(
        document,
        ("FORMA", "SIGNIFICADO", "USO EN LOS DIAGRAMAS"),
        (
            ("Óvalo", "Inicio o fin", "Delimita cada algoritmo."),
            ("Paralelogramo", "Entrada o salida", "Lectura y presentación de resultados."),
            ("Rectángulo", "Proceso", "Cálculo, asignación o llamada modular."),
            ("Rombo", "Decisión", "Validación, condición o control del ciclo."),
        ),
        widths=(1.5, 2.0, 3.5),
        font_size=8.8,
    )
    add_subheading(document, "Criterios técnicos comunes")
    add_bullet(document, "Las entradas inválidas regresan al punto de lectura y no generan resultados parciales.")
    add_bullet(document, "Los ciclos tienen inicialización, condición de terminación y actualización visibles.")
    add_bullet(document, "La variante pública usa información sintética y metadatos genéricos.")

    for exercise in EXERCISES:
        analysis_page = 3 + (exercise.number - 1) * 3
        diagram_page = analysis_page + 1
        code_page = analysis_page + 2

        page_break(document)
        add_page_heading(
            document,
            f"{exercise.number:02d} · Problema {exercise.number} · Análisis",
            exercise.title,
            f"Trazabilidad: análisis p. {analysis_page} · diagrama p. {diagram_page} · pseudocódigo p. {code_page}.",
        )
        add_callout(document, "Enunciado", exercise.statement)
        add_subheading(document, "Entradas, proceso y salida")
        add_table(
            document,
            ("ENTRADAS", "PROCESO", "SALIDA"),
            ((exercise.inputs, exercise.process, exercise.output),),
            widths=(2.15, 2.8, 2.05),
            font_size=8.5,
        )
        add_subheading(document, "Diseño modular")
        add_callout(document, "Módulo", exercise.module, fill=PALE_GOLD, accent="D09A24")
        add_subheading(document, "Variables principales")
        add_table(
            document,
            ("VARIABLE", "TIPO", "FUNCIÓN"),
            exercise.variables,
            widths=(2.0, 1.15, 3.85),
            font_size=8.15,
        )
        add_subheading(document, "Supuestos y restricciones")
        for assumption in exercise.assumptions:
            add_bullet(document, assumption)
        add_subheading(document, "Pruebas de escritorio")
        add_table(
            document,
            ("CASO", "ENTRADA", "RESULTADO ESPERADO"),
            exercise.tests,
            widths=(1.35, 3.2, 2.45),
            font_size=8.1,
        )

        page_break(document)
        add_page_heading(
            document,
            f"{exercise.number:02d} · Problema {exercise.number} · Diseño",
            "Diagrama de flujo",
            "Las rutas Sí/No y las flechas de retorno muestran decisiones y repeticiones.",
        )
        add_picture(document, exercise.diagram_path)
        paragraph = document.add_paragraph()
        format_paragraph(paragraph, after=5, line=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_run(
            paragraph,
            f"Figura {exercise.number}. Flujo de {exercise.title.lower()}.",
            size=8.6,
            italic=True,
            color=MUTED,
        )
        add_callout(
            document,
            "Correspondencia",
            f"El diagrama representa la lectura y los controles aplicables, refleja {exercise.module} y conserva la salida descrita en la tabla IPO.",
            fill=PALE_BLUE,
            accent="4F8397",
        )

        page_break(document)
        add_page_heading(
            document,
            f"{exercise.number:02d} · Problema {exercise.number} · Implementación",
            "Pseudocódigo modular",
            "Notación educativa compatible con las construcciones documentadas por PSeInt.",
        )
        code = exercise.source_path.read_text(encoding="utf-8")
        add_code_block(document, code)

    page_break(document)
    add_page_heading(
        document,
        "11 · Control de calidad",
        "Cobertura de la lista de chequeo",
        "Cada fila corresponde a dos indicadores binarios del instrumento original.",
    )
    coverage_rows = []
    for exercise in EXERCISES:
        analysis_page = 3 + (exercise.number - 1) * 3
        coverage_rows.append(
            (
                f"Problema {exercise.number}",
                f"Diagrama completo — p. {analysis_page + 1}",
                f"Pseudocódigo completo — p. {analysis_page + 2}",
            )
        )
    add_table(
        document,
        ("PROBLEMA", "NOTACIÓN DE DIAGRAMA", "NOTACIÓN DE PSEUDOCÓDIGO"),
        coverage_rows,
        widths=(1.25, 2.85, 2.9),
        font_size=7.9,
    )
    add_subheading(document, "Resultado de las comprobaciones")
    add_bullet(document, "Las implementaciones de referencia ejecutan casos normales, límites y rechazos de entradas inválidas.")
    add_bullet(document, "El ZIP público contiene los diez pseudocódigos, los diez diagramas, el DOCX y un archivo de lectura.")
    add_bullet(document, "El informe personalizado se genera fuera del repositorio y no forma parte de la publicación.")
    add_subheading(document, "Conclusiones")
    add_bullet(document, "La programación modular reduce repeticiones y concentra cada regla de negocio en una función o procedimiento.")
    add_bullet(document, "Las validaciones impiden cálculos indefinidos y hacen explícitos los rangos aceptados.")
    add_bullet(document, "La equivalencia entre tabla IPO, diagrama, pseudocódigo y prueba facilita revisar cada solución.")
    add_subheading(document, "Fuentes")
    add_bullet(document, "Guía de aprendizaje GA3, páginas 5–6: enunciados y lineamientos de entrega.")
    add_bullet(document, "Material oficial Análisis y solución de problemas aplicando algoritmos.")
    add_bullet(document, "Documentación oficial de subprocesos de PSeInt.")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


# ---------------------------------------------------------------------------
# Empaquetado público y variante privada externa
# ---------------------------------------------------------------------------


def add_zip_bytes(archive: zipfile.ZipFile, arcname: str, data: bytes) -> None:
    info = zipfile.ZipInfo(arcname, date_time=(2026, 8, 16, 12, 0, 0))
    info.compress_type = zipfile.ZIP_DEFLATED
    info.external_attr = 0o100644 << 16
    archive.writestr(info, data)


def create_solution_zip(
    destination: Path,
    *,
    document_path: Path,
    root_name: str,
    private_pdf: Path | None = None,
) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    readme = (
        f"{EVIDENCE_CODE}\n"
        "Taller funciones y procedimientos en la solución de algoritmos\n\n"
        "Contenido:\n"
        "- Documento integral en formato editable.\n"
        "- Diez pseudocódigos en notación PSeInt.\n"
        "- Diez diagramas de flujo en formato PNG.\n\n"
        "Todos los casos y nombres de variables son académicos y sintéticos.\n"
    ).encode("utf-8")
    with zipfile.ZipFile(destination, "w") as archive:
        add_zip_bytes(archive, f"{root_name}/LEAME.txt", readme)
        add_zip_bytes(
            archive,
            f"{root_name}/{document_path.name}",
            document_path.read_bytes(),
        )
        if private_pdf is not None:
            add_zip_bytes(
                archive,
                f"{root_name}/{private_pdf.name}",
                private_pdf.read_bytes(),
            )
        for exercise in EXERCISES:
            add_zip_bytes(
                archive,
                f"{root_name}/pseudocodigo/{exercise.source_path.name}",
                exercise.source_path.read_bytes(),
            )
            add_zip_bytes(
                archive,
                f"{root_name}/diagramas/{exercise.diagram_path.name}",
                exercise.diagram_path.read_bytes(),
            )


def export_pdf(source: Path, destination: Path) -> None:
    office = shutil.which("libreoffice") or shutil.which("soffice")
    if not office:
        raise RuntimeError("LibreOffice es obligatorio para generar el PDF privado.")
    destination.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="entrega-privada-") as temp_name:
        temp_dir = Path(temp_name)
        subprocess.run(
            [office, "--headless", "--convert-to", "pdf", "--outdir", str(temp_dir), str(source)],
            check=True,
        )
        generated = temp_dir / f"{source.stem}.pdf"
        if not generated.is_file():
            raise RuntimeError("LibreOffice no creó el PDF privado esperado.")
        shutil.copy2(generated, destination)


def safe_slug(value: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9]+", "_", value.strip())
    return normalized.strip("_") or "Entrega_Privada"


def load_private_profile(path: Path) -> dict[str, str]:
    resolved = path.resolve()
    if resolved.is_relative_to(REPO_ROOT):
        raise ValueError("El perfil privado debe permanecer fuera del repositorio.")
    data = json.loads(path.read_text(encoding="utf-8"))
    name = str(data.get("nombre_completo", "")).strip()
    if not name:
        raise ValueError("El perfil privado requiere nombre_completo.")
    return {str(key): str(value) for key, value in data.items()}


def private_report_text(profile: dict[str, str], docx: Path, pdf: Path, package: Path) -> str:
    optional_lines = []
    labels = {
        "programa": "Programa",
        "institucion": "Institución",
        "centro": "Centro de formación",
        "ciudad": "Ciudad",
        "fecha": "Fecha",
    }
    for field, label in labels.items():
        value = profile.get(field, "").strip()
        if value:
            optional_lines.append(f"- **{label}:** {value}")
    details = "\n".join(optional_lines)
    return f"""# Informe completo local — {EVIDENCE_CODE}

> Documento privado. No publicar ni mover al repositorio.

## Identificación académica

- **Aprendiz:** {profile['nombre_completo']}
{details}
- **Evidencia:** Taller funciones y procedimientos en la solución de algoritmos.
- **Código:** {EVIDENCE_CODE}.

## Alcance resuelto

Se resolvieron los diez problemas establecidos por la guía. Cada problema contiene análisis
de entradas, proceso y salida; función o procedimiento modular; diagrama de flujo;
pseudocódigo y pruebas de escritorio. La entrega cubre los veinte indicadores binarios del
instrumento: una notación gráfica y una notación textual por cada problema.

## Entregables privados

- `{docx.name}`: documento editable personalizado.
- `{pdf.name}`: versión PDF personalizada.
- `{package.name}`: paquete ZIP personalizado para entrega académica.

## Seguridad

Estos archivos se generaron fuera del repositorio público. La variante publicada en GitHub
usa metadatos genéricos y no incorpora los datos de este perfil.
"""


def build_private_delivery(profile_path: Path, output_dir: Path) -> tuple[Path, Path, Path, Path]:
    profile = load_private_profile(profile_path)
    resolved_output = output_dir.resolve()
    if resolved_output.is_relative_to(REPO_ROOT):
        raise ValueError("La salida privada debe permanecer fuera del repositorio.")
    resolved_output.mkdir(parents=True, exist_ok=True)
    person_slug = safe_slug(profile["nombre_completo"])
    stem = f"{person_slug}_{EVIDENCE_CODE}_Taller_Funciones_Procedimientos"
    private_docx = resolved_output / f"{stem}.docx"
    private_pdf = resolved_output / f"{stem}.pdf"
    private_zip = resolved_output / f"{stem}.zip"
    private_report = resolved_output / "INFORME_COMPLETO_PRIVADO.md"

    build_document(private_docx, profile=profile)
    export_pdf(private_docx, private_pdf)
    create_solution_zip(
        private_zip,
        document_path=private_docx,
        private_pdf=private_pdf,
        root_name=stem,
    )
    private_report.write_text(
        private_report_text(profile, private_docx, private_pdf, private_zip),
        encoding="utf-8",
    )
    return private_docx, private_pdf, private_zip, private_report


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--perfil-privado", type=Path)
    parser.add_argument("--salida-privada", type=Path)
    args = parser.parse_args()
    if bool(args.perfil_privado) != bool(args.salida_privada):
        parser.error("--perfil-privado y --salida-privada deben utilizarse juntos.")
    return args


def main() -> None:
    args = parse_args()
    checks = run_tests()
    validate_traceability()
    build_diagrams()
    build_document(PUBLIC_DOCX)
    create_solution_zip(
        PUBLIC_ZIP,
        document_path=PUBLIC_DOCX,
        root_name=EVIDENCE_CODE,
    )
    print(f"Algorithm tests: {checks} passed")
    print(f"Created: {PUBLIC_DOCX}")
    print(f"Created: {PUBLIC_ZIP}")
    if args.perfil_privado:
        for path in build_private_delivery(args.perfil_privado, args.salida_privada):
            print(f"Created private: {path}")


if __name__ == "__main__":
    main()
