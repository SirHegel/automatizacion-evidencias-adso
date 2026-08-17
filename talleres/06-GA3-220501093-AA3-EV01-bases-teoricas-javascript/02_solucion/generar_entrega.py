#!/usr/bin/env python3
"""Genera las ediciones pública y completa local de GA3-220501093-AA3-EV01."""

from __future__ import annotations

import json
import os
from pathlib import Path
import shutil
import subprocess
import tempfile
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOLUTION_DIR = Path(__file__).resolve().parent
WORKSHOP_DIR = SOLUTION_DIR.parent
REPO_ROOT = SOLUTION_DIR.parents[2]
RESOURCE_DIR = SOLUTION_DIR / "recursos" / "figuras"
DELIVERY_DIR = WORKSHOP_DIR / "03_entrega"
PROFILE_PATH = REPO_ROOT / "perfil-aprendiz.local.json"

EVIDENCE_CODE = "GA3-220501093-AA3-EV01"
PUBLIC_STEM = f"{EVIDENCE_CODE}_Bases_Teoricas_JavaScript"
PUBLIC_DOCX = DELIVERY_DIR / f"{PUBLIC_STEM}.docx"
LOCAL_DOCX = DELIVERY_DIR / f"{PUBLIC_STEM}_COMPLETO.local.docx"
LOCAL_PDF = DELIVERY_DIR / f"{PUBLIC_STEM}_COMPLETO.local.pdf"
LOCAL_REPORT = DELIVERY_DIR / "INFORME_COMPLETO.local.md"

FIGURE_PATHS = (
    RESOURCE_DIR / "01_estrategias_ejecucion.png",
    RESOURCE_DIR / "02_capas_javascript.png",
    RESOURCE_DIR / "03_tipos_primitivos.png",
    RESOURCE_DIR / "04_operadores_precedencia.png",
)

GREEN = "39A900"
DARK_GREEN = "174C2C"
DEEP_GREEN = "0E3322"
PALE_GREEN = "EDF7E9"
PALE_BLUE = "EAF3F7"
PALE_GOLD = "FFF4D6"
PALE_RED = "FCECE8"
INK = "18221B"
MUTED = "5E6A62"
LINE = "D8DED9"
WHITE = "FFFFFF"
CODE_BG = "F4F6F4"

FONT_REGULAR_CANDIDATES = (
    Path("/usr/share/fonts/truetype/msttcorefonts/Arial.ttf"),
    Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
)
FONT_BOLD_CANDIDATES = (
    Path("/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf"),
    Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
)


def first_existing(paths: Iterable[Path]) -> Path:
    for path in paths:
        if path.is_file():
            return path
    raise FileNotFoundError("No se encontró una tipografía compatible.")


def image_font(size: int, *, bold: bool = False):
    candidates = FONT_BOLD_CANDIDATES if bold else FONT_REGULAR_CANDIDATES
    return ImageFont.truetype(str(first_existing(candidates)), size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> str:
    output: list[str] = []
    for paragraph in text.splitlines() or [""]:
        words = paragraph.split()
        if not words:
            output.append("")
            continue
        line = words[0]
        for word in words[1:]:
            candidate = f"{line} {word}"
            if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
                line = candidate
            else:
                output.append(line)
                line = word
        output.append(line)
    return "\n".join(output)


def centered_text(
    draw: ImageDraw.ImageDraw,
    bounds: tuple[int, int, int, int],
    text: str,
    *,
    font,
    fill: str = "#18221B",
    padding: int = 28,
) -> None:
    left, top, right, bottom = bounds
    wrapped = wrap_text(draw, text, font, right - left - padding * 2)
    text_box = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=8, align="center")
    width = text_box[2] - text_box[0]
    height = text_box[3] - text_box[1]
    draw.multiline_text(
        ((left + right - width) / 2, (top + bottom - height) / 2 - text_box[1]),
        wrapped,
        font=font,
        fill=fill,
        spacing=8,
        align="center",
    )


def box(
    draw: ImageDraw.ImageDraw,
    bounds: tuple[int, int, int, int],
    text: str,
    *,
    fill: str,
    outline: str = "#174C2C",
    size: int = 28,
) -> None:
    draw.rounded_rectangle(bounds, radius=22, fill=fill, outline=outline, width=4)
    centered_text(draw, bounds, text, font=image_font(size, bold=True))


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line((start, end), fill="#174C2C", width=7)
    x, y = end
    if end[0] > start[0]:
        points = [(x, y), (x - 24, y - 14), (x - 24, y + 14)]
    else:
        points = [(x, y), (x + 24, y - 14), (x + 24, y + 14)]
    draw.polygon(points, fill="#174C2C")


def figure_canvas(title: str, subtitle: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1800, 1000), "#F7FAF7")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((55, 40, 1745, 135), radius=22, fill="#174C2C")
    draw.text((95, 64), title, font=image_font(40, bold=True), fill="#FFFFFF")
    draw.text((90, 160), subtitle, font=image_font(26), fill="#5E6A62")
    return image, draw


def build_execution_figure(path: Path) -> None:
    image, draw = figure_canvas(
        "FIGURA 1 · ESTRATEGIAS DE EJECUCIÓN",
        "Compilar e interpretar describen rutas de trabajo; un motor puede combinarlas.",
    )
    draw.text((95, 235), "COMPILACIÓN AOT", font=image_font(30, bold=True), fill="#174C2C")
    draw.text((95, 560), "JAVASCRIPT MODERNO", font=image_font(30, bold=True), fill="#174C2C")

    top_boxes = (
        ((80, 305, 390, 435), "Código fuente"),
        ((520, 305, 830, 435), "Compilador"),
        ((960, 305, 1270, 435), "Archivo ejecutable"),
        ((1400, 305, 1710, 435), "Procesador"),
    )
    for bounds, label in top_boxes:
        box(draw, bounds, label, fill="#EAF3F7")
    for index in range(len(top_boxes) - 1):
        arrow(draw, (top_boxes[index][0][2] + 12, 370), (top_boxes[index + 1][0][0] - 12, 370))

    bottom_boxes = (
        ((80, 635, 360, 785), "Código\nJavaScript"),
        ((465, 635, 745, 785), "Análisis y\nbytecode"),
        ((850, 635, 1130, 785), "Intérprete\nIgnition"),
        ((1235, 635, 1515, 785), "JIT\nTurboFan"),
    )
    for bounds, label in bottom_boxes:
        box(draw, bounds, label, fill="#EDF7E9")
    for index in range(len(bottom_boxes) - 1):
        arrow(draw, (bottom_boxes[index][0][2] + 12, 710), (bottom_boxes[index + 1][0][0] - 12, 710))
    box(draw, (1380, 830, 1710, 925), "Código nativo optimizado", fill="#FFF4D6", size=25)
    draw.line((1375, 785, 1545, 830), fill="#174C2C", width=6)
    draw.text(
        (92, 900),
        "Idea clave: la implementación elige la estrategia; ECMAScript define el comportamiento del lenguaje.",
        font=image_font(24, bold=True),
        fill="#174C2C",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, optimize=True)


def build_layers_figure(path: Path) -> None:
    image, draw = figure_canvas(
        "FIGURA 2 · LENGUAJE, MOTOR Y ENTORNO",
        "JavaScript no es sinónimo del navegador: cada capa tiene una responsabilidad.",
    )
    layers = (
        ((180, 250, 1620, 400), "ENTORNO ANFITRIÓN\nNavegador · servidor · herramienta", "#EAF3F7"),
        ((300, 445, 1500, 595), "MOTOR DE JAVASCRIPT\nAnaliza, ejecuta, optimiza y administra memoria", "#EDF7E9"),
        ((420, 640, 1380, 790), "ECMASCRIPT\nSintaxis · tipos · operadores · objetos · funciones", "#FFF4D6"),
    )
    for bounds, label, color in layers:
        box(draw, bounds, label, fill=color, size=29)
    draw.text((210, 835), "API del entorno", font=image_font(24, bold=True), fill="#174C2C")
    draw.text((610, 835), "Funciones de primera clase", font=image_font(24, bold=True), fill="#174C2C")
    draw.text((1115, 835), "Objetos con prototipos", font=image_font(24, bold=True), fill="#174C2C")
    draw.text(
        (220, 895),
        "El DOM o el sistema de archivos pertenecen al entorno; no forman parte del núcleo ECMAScript.",
        font=image_font(25),
        fill="#5E6A62",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, optimize=True)


def build_primitives_figure(path: Path) -> None:
    image, draw = figure_canvas(
        "FIGURA 3 · LOS SIETE PRIMITIVOS",
        "Son valores inmutables; Object constituye la categoría no primitiva.",
    )
    cards = (
        ("String", '"texto"', "#EAF3F7"),
        ("Number", "42 · 3.14 · NaN", "#EDF7E9"),
        ("BigInt", "9007199254740993n", "#FFF4D6"),
        ("Boolean", "true · false", "#FCECE8"),
        ("Undefined", "sin asignación", "#EAF3F7"),
        ("Symbol", "identificador único", "#EDF7E9"),
        ("Null", "ausencia intencional", "#FFF4D6"),
    )
    positions = (
        (80, 260, 500, 405),
        (530, 260, 950, 405),
        (980, 260, 1400, 405),
        (305, 455, 725, 600),
        (755, 455, 1175, 600),
        (1205, 455, 1625, 600),
        (530, 650, 950, 795),
    )
    for (name, example, color), bounds in zip(cards, positions):
        box(draw, bounds, f"{name}\n{example}", fill=color, size=27)
    box(draw, (1060, 665, 1625, 805), "Object\nArreglos, funciones y objetos", fill="#FFFFFF", outline="#A84232", size=27)
    draw.text((1080, 825), "NO ES PRIMITIVO", font=image_font(24, bold=True), fill="#A84232")
    draw.text(
        (90, 890),
        'Atención: typeof null produce "object" por compatibilidad histórica; para detectarlo se usa valor === null.',
        font=image_font(25, bold=True),
        fill="#174C2C",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, optimize=True)


def build_operators_figure(path: Path) -> None:
    image, draw = figure_canvas(
        "FIGURA 4 · OPERADORES Y ORDEN DE EVALUACIÓN",
        "La precedencia organiza la expresión; los paréntesis hacen explícita la intención.",
    )
    levels = (
        ((100, 260, 530, 385), "1 · AGRUPACIÓN\n( )", "#EAF3F7"),
        ((685, 260, 1115, 385), "2 · EXPONENTE\n**", "#EDF7E9"),
        ((1270, 260, 1700, 385), "3 · PRODUCTO\n* / %", "#FFF4D6"),
        ((390, 455, 820, 580), "4 · SUMA\n+ −", "#FCECE8"),
        ((980, 455, 1410, 580), "5 · COMPARACIÓN\n<  >  ===", "#EAF3F7"),
    )
    for bounds, label, color in levels:
        box(draw, bounds, label, fill=color, size=28)
    arrow(draw, (535, 323), (675, 323))
    arrow(draw, (1120, 323), (1260, 323))
    draw.line((1485, 385, 1485, 425, 605, 425), fill="#174C2C", width=6)
    arrow(draw, (605, 425), (605, 445))
    arrow(draw, (830, 517), (970, 517))
    box(draw, (155, 675, 775, 825), '2 + 3 * 4  →  14\n(2 + 3) * 4  →  20', fill="#EDF7E9", size=30)
    box(draw, (1025, 675, 1645, 825), '0 == false  →  true\n0 === false  →  false', fill="#FFF4D6", size=30)
    draw.text(
        (175, 880),
        "Recomendación: igualdad estricta, paréntesis cuando aclaren la regla y cortocircuito usado conscientemente.",
        font=image_font(25, bold=True),
        fill="#174C2C",
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, optimize=True)


def build_figures() -> None:
    build_execution_figure(FIGURE_PATHS[0])
    build_layers_figure(FIGURE_PATHS[1])
    build_primitives_figure(FIGURE_PATHS[2])
    build_operators_figure(FIGURE_PATHS[3])


def set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, *, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    properties.append(tag)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(
    document: Document,
    headers: tuple[str, ...],
    rows: Iterable[tuple[str, ...]],
    *,
    widths: tuple[float, ...],
    size: float = 8.3,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, DARK_GREEN)
        set_cell_text(cell, header, bold=True, color=WHITE, size=size)
        cell.width = Inches(widths[index])
    for row_number, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            if row_number % 2 == 1:
                set_cell_shading(cells[index], "F7FAF7")
            set_cell_text(cells[index], value, size=size)
            cells[index].width = Inches(widths[index])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PÁGINA ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.bold = True
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure_document(document: Document, author: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.25)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_after = Pt(5)

    properties = document.core_properties
    properties.author = author
    properties.last_modified_by = author
    properties.title = "Bases teóricas de estructuras de almacenamiento en memoria"
    properties.subject = EVIDENCE_CODE
    properties.keywords = "JavaScript, ECMAScript, tipos primitivos, operadores"

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("SENA · EVIDENCIA DE CONOCIMIENTO")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    code = header.add_run(f"                                                        {EVIDENCE_CODE}")
    code.font.name = "Arial"
    code.font.size = Pt(8)
    code.font.bold = True
    code.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    add_page_number(section.footer.paragraphs[0])


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: float = 10, color: str = INK):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_body(document: Document, text: str, *, size: float = 9.7) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.paragraph_format.space_after = Pt(6)
    add_run(paragraph, text, size=size)


def add_bullets(document: Document, items: Iterable[str], *, size: float = 9.3) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.05
        add_run(paragraph, item, size=size)


def add_heading(document: Document, kicker: str, title: str, subtitle: str = "") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    add_run(paragraph, kicker.upper(), bold=True, size=8.5, color=GREEN)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    add_run(paragraph, title, bold=True, size=23, color=DARK_GREEN)
    if subtitle:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        add_run(paragraph, subtitle, italic=True, size=9.3, color=MUTED)


def add_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    add_run(paragraph, text, bold=True, size=12.2, color=DARK_GREEN)


def add_callout(document: Document, title: str, text: str, *, fill: str = PALE_GREEN, accent: str = GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=100, start=130, bottom=100, end=130)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    add_run(paragraph, title.upper(), bold=True, size=8, color=accent)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    add_run(paragraph, text, size=9.2)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    set_cell_margins(cell, top=90, start=130, bottom=90, end=130)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(code)
    run.font.name = "DejaVu Sans Mono"
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor.from_string(DEEP_GREEN)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(document: Document, path: Path, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(path), width=Inches(7.0))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    add_run(paragraph, caption, italic=True, size=8.2, color=MUTED)


def page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_cover(document: Document, profile: dict[str, str] | None) -> None:
    for _ in range(2):
        document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, "EVIDENCIA DE CONOCIMIENTO", bold=True, size=10, color=GREEN)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(28)
    paragraph.paragraph_format.space_after = Pt(10)
    add_run(paragraph, "Bases teóricas de estructuras\nde almacenamiento en memoria", bold=True, size=27, color=DARK_GREEN)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, EVIDENCE_CODE, bold=True, size=11, color=MUTED)

    document.add_paragraph()
    if profile:
        rows = (
            ("Presentado por", profile["nombre_completo"]),
            (profile.get("tipo_documento", "Documento"), profile["documento"]),
            ("Programa", profile["programa"]),
            ("Institución", profile["institucion"]),
            ("Fecha", profile["fecha"]),
        )
        label = "VERSIÓN COMPLETA LOCAL"
    else:
        rows = (
            ("Programa", "Análisis y Desarrollo de Software"),
            ("Institución", "Servicio Nacional de Aprendizaje SENA"),
            ("Fase", "Planeación"),
            ("Naturaleza", "Documento académico público"),
        )
        label = "VERSIÓN PÚBLICA SIN DATOS PERSONALES"
    add_table(document, ("CAMPO", "INFORMACIÓN"), rows, widths=(1.7, 5.1), size=9.2)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    add_run(paragraph, label, bold=True, size=9.5, color=GREEN)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    add_run(
        paragraph,
        "Lenguajes · JavaScript · Tipos primitivos · Operadores",
        size=10,
        color=MUTED,
    )


def build_document(output: Path, profile: dict[str, str] | None = None) -> None:
    author = profile["nombre_completo"] if profile else "Entrega académica pública"
    document = Document()
    configure_document(document, author)
    add_cover(document, profile)

    page_break(document)
    add_heading(document, "00 · Punto de partida", "Propósito y ruta de lectura", "Los cinco indicadores se atienden de forma explícita y verificable.")
    add_callout(
        document,
        "Propósito",
        "Consolidar las bases necesarias para leer y escribir JavaScript con criterio: comprender cómo se ejecuta un lenguaje, qué caracteriza a ECMAScript, cómo se representan los valores simples y cómo los operadores transforman esos valores.",
    )
    add_subheading(document, "Objetivos específicos")
    add_bullets(
        document,
        (
            "Comparar compilación e interpretación sin presentar categorías absolutas.",
            "Distinguir el lenguaje JavaScript del motor y de las API ofrecidas por cada entorno.",
            "Reconocer los siete primitivos vigentes y seleccionar su uso de acuerdo con el dato.",
            "Aplicar operadores atendiendo a precedencia, coerción y evaluación por cortocircuito.",
        ),
    )
    add_subheading(document, "Ruta del documento")
    add_table(
        document,
        ("CAPÍTULO", "PREGUNTA QUE RESUELVE", "APOYO VISUAL"),
        (
            ("1", "¿Cómo llega el código fuente a ejecutarse?", "Dos canales de ejecución."),
            ("2", "¿Qué hace distintivo a JavaScript?", "Capas lenguaje–motor–entorno."),
            ("3", "¿Qué valores primitivos existen y para qué sirven?", "Mapa de siete tipos."),
            ("4", "¿Cómo se forman y evalúan expresiones?", "Escalera de precedencia."),
        ),
        widths=(0.8, 3.9, 2.1),
    )
    add_body(document, "Método: síntesis comparativa de la guía SENA [1], la especificación ECMA-262 [2], MDN [3–6] y la documentación del motor V8 [7]. Las figuras son elaboración propia y declaran su base conceptual.", size=9.2)

    page_break(document)
    add_heading(document, "01 · Estrategias de ejecución", "Lenguajes compilados e interpretados", "La diferencia principal está en cuándo y cómo se traduce el programa.")
    add_body(document, "Compilar significa transformar un programa a otro formato antes o durante su ejecución. En un flujo AOT (ahead of time), la traducción suele producir un archivo ejecutable antes de iniciar el programa. Interpretar significa que un programa intermediario analiza y ejecuta instrucciones durante la marcha. Son estrategias de implementación: un lenguaje no queda obligado para siempre a una sola de ellas [5].")
    add_table(
        document,
        ("ASPECTO", "COMPILACIÓN AOT", "INTERPRETACIÓN"),
        (
            ("Momento", "Traducción previa a la ejecución.", "Análisis y ejecución durante el uso."),
            ("Producto", "Puede generar binario o código objeto.", "Normalmente conserva el fuente y usa un intérprete."),
            ("Errores", "Muchos se detectan al compilar.", "Algunos aparecen al alcanzar la instrucción."),
            ("Inicio", "Requiere una fase previa de construcción.", "Puede iniciar sin crear un ejecutable independiente."),
            ("Rendimiento", "El ejecutable ya está traducido.", "La traducción añade trabajo en tiempo de ejecución."),
            ("Portabilidad", "El binario depende del destino.", "El fuente viaja si existe un entorno compatible."),
        ),
        widths=(1.15, 2.85, 2.85),
        size=8.1,
    )
    add_callout(document, "Conclusión comparativa", "Compilado no significa automáticamente mejor, ni interpretado significa necesariamente lento. Importan el motor, la optimización, la carga de trabajo y el momento en que se realiza la traducción.", fill=PALE_GOLD, accent="D09A24")

    page_break(document)
    add_heading(document, "01 · Estrategias de ejecución", "El caso real de JavaScript", "Los motores modernos combinan intérprete y compilador JIT.")
    add_figure(document, FIGURE_PATHS[0], "Figura 1. Estrategias de ejecución. Elaboración propia con base en MDN [5] y V8 [7].")
    add_body(document, "ECMAScript especifica la sintaxis y los resultados observables, pero no obliga a una arquitectura interna. En V8, el código pasa por análisis y bytecode, se ejecuta con Ignition y las partes usadas con frecuencia pueden optimizarse con TurboFan. Por eso es más preciso afirmar que JavaScript suele ejecutarse mediante una combinación de interpretación y compilación JIT [7].")
    add_callout(document, "AOT frente a JIT", "AOT traduce antes de iniciar; JIT compila durante la ejecución usando información real del programa. Una misma aplicación puede recorrer bytecode y código nativo optimizado.", fill=PALE_BLUE, accent="4F8397")

    page_break(document)
    add_heading(document, "02 · Lenguaje", "Características principales de JavaScript", "JavaScript está estandarizado mediante la especificación ECMAScript.")
    add_table(
        document,
        ("CARACTERÍSTICA", "SIGNIFICADO", "IMPLICACIÓN"),
        (
            ("Estandarizado", "ECMA-262 define la semántica.", "Los motores comparten una base común."),
            ("Dinámico", "La variable no queda unida a un tipo.", "El tipo pertenece al valor y puede cambiar."),
            ("Coerción", "Admite conversiones implícitas.", "Conviene controlar comparaciones y entradas."),
            ("Multiparadigma", "Combina estilos imperativo, funcional y orientado a objetos.", "Permite escoger una solución adecuada al problema."),
            ("Prototipos", "Los objetos delegan mediante cadenas de prototipos.", "La herencia no depende solo de clases."),
            ("Primera clase", "Las funciones son valores invocables.", "Pueden pasarse, retornarse y almacenarse."),
            ("Memoria automática", "El motor recupera objetos no alcanzables.", "No se libera memoria manualmente."),
            ("Portable", "Funciona en distintos anfitriones.", "La API disponible depende del entorno."),
        ),
        widths=(1.35, 2.75, 2.75),
        size=7.75,
    )
    add_callout(document, "Precisión terminológica", "JavaScript es el lenguaje; ECMAScript es su especificación; V8, SpiderMonkey o JavaScriptCore son motores; navegador y Node.js son entornos que añaden API.", fill=PALE_GOLD, accent="D09A24")

    page_break(document)
    add_heading(document, "02 · Lenguaje", "Tres capas que conviene separar", "Esta distinción evita atribuir al lenguaje funciones que pertenecen al entorno.")
    add_figure(document, FIGURE_PATHS[1], "Figura 2. Capas de una ejecución JavaScript. Elaboración propia con base en ECMA-262 [2] y MDN [3].")
    add_body(document, "El núcleo ECMAScript aporta valores, operadores, control, objetos, funciones, clases y módulos. El motor materializa esas reglas y se ocupa de ejecutar y optimizar. El anfitrión agrega capacidades: un navegador ofrece el DOM y eventos; un entorno de servidor puede ofrecer sistema de archivos o red. Esas API no forman parte del lenguaje [2, 3].")
    add_code(document, 'const duplicar = (valor) => valor * 2;\nconst operaciones = [duplicar];\nconsole.log(operaciones[0](21)); // 42')
    add_body(document, "El ejemplo muestra una función como valor: se guarda dentro de un arreglo y después se invoca. Esa propiedad sustenta callbacks, composición y gran parte de la programación asíncrona.", size=9.2)

    page_break(document)
    add_heading(document, "03 · Valores", "Tipos de datos primitivos", "ECMA-262 define siete tipos primitivos y Object como categoría separada.")
    add_table(
        document,
        ("TIPO", "EJEMPLO", "USO HABITUAL", "typeof"),
        (
            ("String", '"Bogotá"', "Texto, etiquetas y entradas.", '"string"'),
            ("Number", "42; 3.14; NaN", "Cálculo general con enteros o decimales.", '"number"'),
            ("BigInt", "9007199254740993n", "Enteros fuera del rango seguro de Number.", '"bigint"'),
            ("Boolean", "true; false", "Decisiones y estados lógicos.", '"boolean"'),
            ("Undefined", "let dato;", "Ausencia de asignación o retorno.", '"undefined"'),
            ("Symbol", 'Symbol("clave")', "Identificadores únicos de propiedades.", '"symbol"'),
            ("Null", "null", "Ausencia intencional de un valor de objeto.", '"object"*'),
        ),
        widths=(1.0, 1.7, 3.05, 1.05),
        size=7.75,
    )
    add_body(document, "Un primitivo es inmutable: la operación crea otro valor en vez de modificar el existente. La variable sí puede reasignarse. Arreglos, funciones, fechas y objetos pertenecen a Object, por lo que no deben sumarse como un octavo primitivo [2, 4].")
    add_callout(document, "BigInt y Number", "No deben mezclarse directamente en aritmética. BigInt conserva enteros arbitrariamente grandes, pero no representa fracciones. Number usa punto flotante y tiene límites de precisión.", fill=PALE_BLUE, accent="4F8397")

    page_break(document)
    add_heading(document, "03 · Valores", "Uso, detección y dos excepciones", "El tipo correcto comunica intención y evita conversiones inesperadas.")
    add_figure(document, FIGURE_PATHS[2], "Figura 3. Mapa de tipos primitivos. Elaboración propia con base en ECMA-262 [2] y MDN [4].")
    add_code(document, 'const titulo = "Informe";          // String\nconst paginas = 12;                 // Number\nconst publicado = true;             // Boolean\nlet observacion;                     // Undefined\nconst seleccion = null;              // Null\nconst clave = Symbol("registro");    // Symbol')
    add_bullets(
        document,
        (
            '`typeof null` devuelve `"object"` por compatibilidad histórica; la comprobación segura es `valor === null`.',
            "NaN pertenece al tipo Number, pero representa un resultado numérico inválido; se detecta con Number.isNaN(valor).",
            "Una propiedad inexistente suele producir undefined; null suele expresar una ausencia decidida por el programa.",
        ),
        size=8.9,
    )

    page_break(document)
    add_heading(document, "04 · Expresiones", "Operadores en JavaScript", "Un operador recibe operandos y produce un valor o un efecto.")
    add_table(
        document,
        ("FAMILIA", "OPERADORES", "EJEMPLO Y RESULTADO"),
        (
            ("Aritméticos", "+  −  *  /  %  **", "2 ** 3 → 8"),
            ("Asignación", "=  +=  -=  *=  ??=", "total += 5"),
            ("Comparación", "<  >  <=  >=", "edad >= 18 → booleano"),
            ("Igualdad", "===  !==  ==  !=", '5 === "5" → false'),
            ("Lógicos", "&&  ||  !", "activo && autorizado"),
            ("Fusión nula", "??", 'nombre ?? "Sin dato"'),
            ("Condicional", "? :", 'mayor ? "Sí" : "No"'),
            ("Unarios", "typeof  delete  void  ++", "typeof 42 → number"),
            ("Relacionales", "in  instanceof", '"id" in objeto'),
            ("Bit a bit", "&  |  ^  ~  <<  >>  >>>", "5 & 1 → 1"),
        ),
        widths=(1.25, 2.1, 3.55),
        size=7.7,
    )
    add_callout(document, "Igualdad estricta", "`===` y `!==` comparan sin convertir tipos y suelen expresar mejor la intención. La igualdad flexible puede ser válida, pero exige conocer con precisión sus reglas de coerción [6].", fill=PALE_GOLD, accent="D09A24")

    page_break(document)
    add_heading(document, "04 · Expresiones", "Precedencia, coerción y cortocircuito", "Los paréntesis y la igualdad estricta vuelven el resultado más predecible.")
    add_figure(document, FIGURE_PATHS[3], "Figura 4. Orden de evaluación. Elaboración propia con base en MDN [6].")
    add_code(document, 'const subtotal = 85000;\nconst descuento = subtotal > 50000 ? subtotal * 0.20 : 0;\nconst total = subtotal - descuento;       // 68000\n\nconst alias = null;\nconst etiqueta = alias ?? "Sin registrar";')
    add_bullets(
        document,
        (
            "La multiplicación se evalúa antes que la suma; los paréntesis pueden cambiar o aclarar ese orden.",
            '`"5" + 2` produce `"52"`; `Number("5") + 2` produce `7`. La conversión explícita comunica la intención.',
            "`&&`, `||` y `??` usan cortocircuito: el segundo operando se evalúa solo cuando la regla lo necesita.",
        ),
        size=8.8,
    )

    page_break(document)
    add_heading(document, "05 · Cierre", "Conclusiones y matriz de cumplimiento", "El conocimiento conceptual se traduce en decisiones de código más seguras.")
    add_bullets(
        document,
        (
            "Compilación e interpretación son estrategias combinables; JavaScript moderno usa ejecución adaptativa.",
            "La especificación, el motor y el entorno son capas distintas y complementarias.",
            "Los siete primitivos representan valores simples e inmutables; Object agrupa las estructuras compuestas.",
            "Los operadores requieren considerar tipo, precedencia, coerción y cortocircuito.",
        ),
    )
    add_subheading(document, "Cobertura de la lista de chequeo")
    add_table(
        document,
        ("N.º", "INDICADOR", "EVIDENCIA", "ESTADO"),
        (
            ("1", "Compilados e interpretados", "Páginas 3–4, tabla y figura 1.", "Cumple"),
            ("2", "Características de JavaScript", "Páginas 5–6, tabla, figura 2 y código.", "Cumple"),
            ("3", "Tipos primitivos y uso", "Páginas 7–8, tabla, figura 3 y ejemplos.", "Cumple"),
            ("4", "Operadores", "Páginas 9–10, tabla, figura 4 y ejemplos.", "Cumple"),
            ("5", "Imágenes y fuentes", "Cuatro figuras propias y referencias numeradas.", "Cumple"),
        ),
        widths=(0.45, 2.3, 3.25, 0.8),
        size=7.8,
    )
    add_callout(document, "Resultado", "Los cinco indicadores tienen evidencia directa. El documento evita extender el alcance hacia los ejercicios de programación que corresponden a la evidencia posterior.")

    page_break(document)
    add_heading(document, "06 · Fuentes", "Referencias", "Consultadas el 16 de agosto de 2026. Enlaces directos a la fuente utilizada.")
    references = (
        "[1] Servicio Nacional de Aprendizaje SENA. (s. f.). Guía de aprendizaje: fase de planeación, actividad GA3-220501093-AA3, pp. 8–9. https://archivos.territorio.la/archivos/clases/Guianaprendizajen3___58631be32843215___.pdf",
        "[2] Ecma International. (2026). ECMAScript Language Specification (ECMA-262). https://tc39.es/ecma262/",
        "[3] MDN Web Docs. (2025). JavaScript language overview. https://developer.mozilla.org/en-US/docs/Web/JavaScript/Guide/Language_overview",
        "[4] MDN Web Docs. (2025). JavaScript data types and data structures. https://developer.mozilla.org/en-US/docs/Web/JavaScript/Guide/Data_structures",
        "[5] MDN Web Docs. (2025). Compile. https://developer.mozilla.org/en-US/docs/Glossary/Compile",
        "[6] MDN Web Docs. (2025). Expressions and operators. https://developer.mozilla.org/en-US/docs/Web/JavaScript/Guide/Expressions_and_operators",
        "[7] V8 Project. (2017). Launching Ignition and TurboFan. https://v8.dev/blog/launching-ignition-and-turbofan",
    )
    for reference in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.24)
        paragraph.paragraph_format.first_line_indent = Inches(-0.24)
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.line_spacing = 1.05
        add_run(paragraph, reference, size=8.7)
    add_subheading(document, "Procedencia de las ilustraciones")
    add_body(document, "Las figuras 1–4 son elaboración propia generada específicamente para esta evidencia. No contienen imágenes, logotipos ni capturas de terceros. Sus pies señalan las fuentes conceptuales que sustentan cada representación.", size=9.1)
    add_callout(document, "Control de versión", "La edición pública y la edición completa local comparten exactamente este desarrollo académico. Únicamente cambian la portada y los metadatos del documento.", fill=PALE_BLUE, accent="4F8397")

    atomic_save_document(document, output)


def atomic_save_document(document: Document, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    descriptor, temp_name = tempfile.mkstemp(prefix=f".{destination.stem}-", suffix=destination.suffix, dir=destination.parent)
    os.close(descriptor)
    temporary = Path(temp_name)
    try:
        document.save(temporary)
        os.replace(temporary, destination)
    finally:
        temporary.unlink(missing_ok=True)


def export_pdf(source: Path, destination: Path) -> None:
    libreoffice = shutil.which("libreoffice")
    if not libreoffice:
        raise RuntimeError("LibreOffice es obligatorio para crear el PDF completo local.")
    with tempfile.TemporaryDirectory(prefix="ev01-local-") as temp_directory:
        output_directory = Path(temp_directory)
        subprocess.run(
            [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_directory), str(source)],
            check=True,
        )
        generated = output_directory / f"{source.stem}.pdf"
        if not generated.is_file():
            raise RuntimeError(f"LibreOffice no produjo el PDF esperado para {source}.")
        destination.parent.mkdir(parents=True, exist_ok=True)
        descriptor, temp_name = tempfile.mkstemp(
            prefix=f".{destination.stem}-",
            suffix=destination.suffix,
            dir=destination.parent,
        )
        os.close(descriptor)
        temporary = Path(temp_name)
        try:
            shutil.copy2(generated, temporary)
            os.replace(temporary, destination)
        finally:
            temporary.unlink(missing_ok=True)


def load_profile() -> dict[str, str] | None:
    if not PROFILE_PATH.is_file():
        return None
    profile = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    required = ("nombre_completo", "documento", "programa", "institucion", "fecha")
    missing = [field for field in required if not str(profile.get(field, "")).strip()]
    if missing:
        raise ValueError(f"El perfil local no contiene los campos requeridos: {missing}")
    return {key: str(value).strip() for key, value in profile.items()}


def ensure_local_destination(path: Path) -> None:
    resolved = path.resolve()
    delivery = DELIVERY_DIR.resolve()
    if resolved.parent != delivery or ".local." not in path.name:
        raise ValueError(f"La entrega completa debe usar un nombre .local dentro de {DELIVERY_DIR}.")
    relative = path.relative_to(REPO_ROOT)
    ignored = subprocess.run(
        ["git", "check-ignore", "--no-index", "-q", "--", relative.as_posix()],
        cwd=REPO_ROOT,
        check=False,
    )
    if ignored.returncode != 0:
        raise ValueError(f"Git no está ignorando el archivo completo local: {relative}")
    tracked = subprocess.run(
        ["git", "ls-files", "--error-unmatch", "--", relative.as_posix()],
        cwd=REPO_ROOT,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        check=False,
    )
    if tracked.returncode == 0:
        raise ValueError(f"El archivo completo local ya está rastreado por Git: {relative}")


def write_local_report(profile: dict[str, str]) -> None:
    text = f"""# Informe completo local — {EVIDENCE_CODE}

- Presentado por: {profile['nombre_completo']}
- {profile.get('tipo_documento', 'Documento')}: {profile['documento']}
- Programa: {profile['programa']}
- Institución: {profile['institucion']}
- Fecha: {profile['fecha']}

## Entregables

- `{LOCAL_DOCX.name}`
- `{LOCAL_PDF.name}`

La solución cubre los cinco indicadores del instrumento y contiene cuatro ilustraciones
originales con fuentes. Este informe y los dos entregables tienen uso exclusivamente local;
la automatización impide que una ruta `.local` sea rastreada por Git.
"""
    descriptor, temp_name = tempfile.mkstemp(prefix=".informe-local-", suffix=".md", dir=DELIVERY_DIR)
    os.close(descriptor)
    temporary = Path(temp_name)
    try:
        temporary.write_text(text, encoding="utf-8")
        os.replace(temporary, LOCAL_REPORT)
    finally:
        temporary.unlink(missing_ok=True)


def build_local_delivery(profile: dict[str, str]) -> None:
    for path in (LOCAL_DOCX, LOCAL_PDF, LOCAL_REPORT):
        ensure_local_destination(path)
    build_document(LOCAL_DOCX, profile)
    export_pdf(LOCAL_DOCX, LOCAL_PDF)
    write_local_report(profile)
    for path in (PROFILE_PATH, LOCAL_DOCX, LOCAL_PDF, LOCAL_REPORT):
        path.chmod(0o600)


def main() -> None:
    build_figures()
    build_document(PUBLIC_DOCX)
    print(f"Created public: {PUBLIC_DOCX}")
    profile = load_profile()
    if profile:
        build_local_delivery(profile)
        print("Created complete local edition beside the public delivery.")
    else:
        print("Local profile absent; only the public edition was generated.")


if __name__ == "__main__":
    main()
