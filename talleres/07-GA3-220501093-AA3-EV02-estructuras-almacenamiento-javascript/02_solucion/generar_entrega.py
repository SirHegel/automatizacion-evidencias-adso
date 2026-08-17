#!/usr/bin/env python3
"""Valida las soluciones y genera los PDF público y personalizado de la EV02."""

from __future__ import annotations

import argparse
from dataclasses import dataclass
import json
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_PATH = Path(__file__).resolve()
SOLUTION_DIR = SCRIPT_PATH.parent
WORKSHOP_DIR = SOLUTION_DIR.parent
REPO_ROOT = WORKSHOP_DIR.parents[1]
CODE = "GA3-220501093-AA3-EV02"
PROFILE_PATH = REPO_ROOT / "perfil-aprendiz.local.json"
PUBLIC_DIR = WORKSHOP_DIR / "03_entrega"
PUBLIC_PDF = PUBLIC_DIR / f"{CODE}_Resolucion_Estructuras_Almacenamiento.pdf"
PERSONAL_DIR = WORKSHOP_DIR / "04_entrega_personalizada.local"
TEST_PATH = SOLUTION_DIR / "pruebas" / "soluciones.test.js"

GREEN = "39A900"
DARK_GREEN = "174C2C"
PALE_GREEN = "EDF7E9"
PALE_BLUE = "EAF3F7"
INK = "18221B"
MUTED = "5E6A62"
LINE = "D8DED9"
CODE_BG = "F4F6F4"
WHITE = "FFFFFF"
PUBLIC_AUTHOR = "Entrega académica pública"


@dataclass(frozen=True)
class Problem:
    number: int
    title: str
    source: Path
    objective: str
    strategy: str
    example: str
    validations: tuple[str, ...]
    tests: str


PROBLEMS = (
    Problem(
        1,
        "Perímetros y áreas de figuras planas",
        SOLUTION_DIR / "codigo" / "01_figuras_planas.js",
        "Calcular perímetro y área de triángulo, rectángulo, cuadrado y círculo.",
        "Se usa un despachador por tipo de figura y una función independiente para cada fórmula. "
        "La circunferencia se calcula correctamente como 2 × π × radio.",
        "Con radio 3, el círculo produce un perímetro aproximado de 18,8496 y un área de 28,2743.",
        (
            "Todas las medidas deben ser numéricas, finitas y mayores que cero.",
            "Los tres lados del triángulo deben cumplir la desigualdad triangular.",
            "Solo se aceptan las cuatro figuras contempladas por la guía.",
        ),
        "Tres pruebas cubren cálculos, despacho, formato y dominios inválidos.",
    ),
    Problem(
        2,
        "Análisis de diez edades",
        SOLUTION_DIR / "codigo" / "02_analisis_edades.js",
        "Leer diez edades y obtener menores, adultos, adultos mayores, mínimo, máximo y promedio.",
        "El arreglo se valida antes de calcular las estadísticas en una sola reducción. Las "
        "categorías son excluyentes: 1–17, 18–59 y 60 años o más.",
        "Para [10, 17, 18, 25, 40, 59, 60, 70, 80, 90], el mínimo es 10, el máximo 90 y el promedio 46,9.",
        (
            "El arreglo debe contener exactamente diez elementos.",
            "Cada edad debe ser un entero entre 1 y 120.",
            "El promedio conserva precisión y se presenta con formato legible.",
        ),
        "Dos pruebas verifican estadísticas, límites y rechazo de edades inválidas.",
    ),
    Problem(
        3,
        "Mezcla ordenada de dos vectores",
        SOLUTION_DIR / "codigo" / "03_mezclar_vectores.js",
        "Combinar dos vectores ascendentes, de máximo cinco elementos cada uno, sin perder el orden.",
        "Se aplica el algoritmo de mezcla con dos índices. Su costo es lineal O(n + m), conserva "
        "duplicados y evita ordenar de nuevo el resultado.",
        "La mezcla de [1, 4, 7] y [2, 4, 9] genera [1, 2, 4, 4, 7, 9].",
        (
            "Cada vector debe tener entre uno y cinco enteros.",
            "Los valores de entrada deben encontrarse en orden ascendente.",
            "La función devuelve un arreglo nuevo y no modifica las entradas.",
        ),
        "Dos pruebas cubren mezcla, duplicados, tamaños, tipos y orden inválido.",
    ),
    Problem(
        4,
        "Encuesta musical persistente",
        SOLUTION_DIR / "codigo" / "04_encuesta_musical.js",
        "Administrar los datos de hasta seis personas y sus canciones favoritas mediante un archivo JSON.",
        "La lógica de dominio permanece separada del menú interactivo y de la persistencia. Se "
        "implementan alta, consulta, listado, modificación, eliminación, guardado y recarga.",
        "Un registro sintético DEMO-001 puede agregarse, consultarse, modificarse y eliminarse sin afectar los demás datos.",
        (
            "La identificación es única y cada campo se normaliza antes de almacenarse.",
            "Cada persona registra entre una y tres canciones con artista y título.",
            "El reemplazo del archivo se realiza de forma atómica y se rechaza JSON dañado.",
        ),
        "Seis pruebas cubren validación, CRUD, límite de registros y persistencia segura.",
    ),
)

REQUIRED_FILES = (
    *(problem.source for problem in PROBLEMS),
    TEST_PATH,
    SOLUTION_DIR / "package.json",
    SOLUTION_DIR / "INSTRUCCIONES.md",
    SOLUTION_DIR / "FUENTES_Y_DECISIONES.md",
)


def run(command: list[str], *, cwd: Path = SOLUTION_DIR) -> subprocess.CompletedProcess[str]:
    print("→ " + " ".join(command))
    return subprocess.run(command, cwd=cwd, check=True, text=True)


def validate_sources() -> None:
    missing = [path for path in REQUIRED_FILES if not path.is_file() or path.stat().st_size == 0]
    if missing:
        raise RuntimeError(f"Faltan fuentes obligatorias: {missing}")
    node = shutil.which("node")
    if not node:
        raise RuntimeError("Node.js no está disponible; no se pueden validar las soluciones.")
    for problem in PROBLEMS:
        run([node, "--check", str(problem.source)])
    run([node, "--test", str(TEST_PATH)])


def filename_fragment(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    fragment = re.sub(r"[^A-Za-z0-9]+", "_", ascii_value).strip("_")
    if not fragment:
        raise RuntimeError("El nombre del perfil no permite construir un nombre de archivo seguro.")
    return fragment


def load_profile() -> dict[str, str] | None:
    if not PROFILE_PATH.is_file():
        return None
    profile = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    required = (
        "nombre_completo",
        "tipo_documento",
        "documento",
        "programa",
        "institucion",
        "fecha",
    )
    missing = [key for key in required if not str(profile.get(key, "")).strip()]
    if missing:
        raise RuntimeError(f"El perfil local no contiene los campos requeridos: {missing}")
    return {key: str(value).strip() for key, value in profile.items()}


def git_ignored(path: Path) -> bool:
    relative = path.relative_to(REPO_ROOT)
    result = subprocess.run(
        ["git", "check-ignore", "--no-index", "-q", "--", relative.as_posix()],
        cwd=REPO_ROOT,
        check=False,
    )
    return result.returncode == 0


def git_tracked(path: Path) -> bool:
    relative = path.relative_to(REPO_ROOT)
    result = subprocess.run(
        ["git", "ls-files", "--error-unmatch", "--", relative.as_posix()],
        cwd=REPO_ROOT,
        check=False,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return result.returncode == 0


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, size: float = 10, color: str = INK, font: str = "Arial"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = add_run(paragraph, "PÁGINA ", bold=True, size=8, color=DARK_GREEN)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure_document(document: Document, author: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    properties = document.core_properties
    properties.author = author
    properties.last_modified_by = author
    properties.title = "Resolución a problemas algorítmicos aplicando estructuras de almacenamiento"
    properties.subject = CODE
    properties.keywords = "JavaScript, arreglos, matrices, archivos, algoritmos"

    header = section.header.paragraphs[0]
    add_run(header, "SENA · EVIDENCIA DE DESEMPEÑO", bold=True, size=8, color=DARK_GREEN)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code_run = header.add_run(f"                                             {CODE}")
    code_run.font.name = "Arial"
    code_run.font.size = Pt(8)
    code_run.font.bold = True
    code_run.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    add_page_number(section.footer.paragraphs[0])


def add_table(document: Document, headers: tuple[str, ...], rows: Iterable[tuple[str, ...]], widths: tuple[float, ...]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, DARK_GREEN)
        set_cell_margins(cell, top=85, start=100, bottom=85, end=100)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_run(paragraph, value, bold=True, size=8, color=WHITE)
    for row_number, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].width = Inches(widths[index])
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index], top=75, start=95, bottom=75, end=95)
            if row_number % 2:
                set_cell_shading(cells[index], "F7FAF7")
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_run(paragraph, value, size=8.2)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(document: Document, label: str, title: str, subtitle: str = "") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    add_run(paragraph, label.upper(), bold=True, size=8.2, color=GREEN)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    add_run(paragraph, title, bold=True, size=21, color=DARK_GREEN)
    if subtitle:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(7)
        add_run(paragraph, subtitle, italic=True, size=9, color=MUTED)


def add_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    add_run(paragraph, text, bold=True, size=11.5, color=DARK_GREEN)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.paragraph_format.space_after = Pt(6)
    add_run(paragraph, text, size=9.3)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        paragraph.paragraph_format.space_after = Pt(3)
        add_run(paragraph, item, size=9)


def add_callout(document: Document, title: str, text: str, fill: str = PALE_GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=100, start=130, bottom=100, end=130)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    add_run(paragraph, title.upper(), bold=True, size=8, color=GREEN)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    add_run(paragraph, text, size=9)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(document: Document, source: Path) -> None:
    code = source.read_text(encoding="utf-8").rstrip()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(7.8)
    paragraph.paragraph_format.keep_together = False
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_BG)
    properties.append(shading)
    add_run(paragraph, code, size=7.5, color=INK, font="DejaVu Sans Mono")


def page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_cover(document: Document, profile: dict[str, str] | None) -> None:
    for _ in range(2):
        document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, "EVIDENCIA DE DESEMPEÑO", bold=True, size=10, color=GREEN)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(25)
    paragraph.paragraph_format.space_after = Pt(9)
    add_run(
        paragraph,
        "Resolución a problemas algorítmicos aplicando estructuras de almacenamiento",
        bold=True,
        size=24,
        color=DARK_GREEN,
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, CODE, bold=True, size=11, color=MUTED)
    document.add_paragraph()

    if profile:
        rows = (
            ("Presentado por", profile["nombre_completo"]),
            (profile["tipo_documento"], profile["documento"]),
            ("Programa", profile["programa"]),
            ("Institución", profile["institucion"]),
            ("Fecha", profile["fecha"]),
        )
        label = "ENTREGA PERSONALIZADA PARA SENA"
    else:
        rows = (
            ("Programa", "Análisis y Desarrollo de Software"),
            ("Institución", "Servicio Nacional de Aprendizaje SENA"),
            ("Fase", "Planeación"),
            ("Naturaleza", "Documento académico público"),
        )
        label = "VERSIÓN PÚBLICA SIN DATOS PERSONALES"
    add_table(document, ("CAMPO", "INFORMACIÓN"), rows, (1.7, 5.1))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    add_run(paragraph, label, bold=True, size=9, color=GREEN)


def build_document(destination: Path, profile: dict[str, str] | None = None) -> None:
    document = Document()
    configure_document(document, profile["nombre_completo"] if profile else PUBLIC_AUTHOR)
    add_cover(document, profile)

    page_break(document)
    add_heading(
        document,
        "00 · Presentación",
        "Alcance, método y evidencias",
        "Cuatro soluciones ejecutables en JavaScript, acompañadas de análisis y pruebas.",
    )
    add_callout(
        document,
        "Resultado",
        "Los cuatro indicadores del instrumento se atienden mediante programas independientes, "
        "validaciones explícitas, funciones reutilizables y trece pruebas automatizadas.",
    )
    add_body(
        document,
        "Resolución a problemas algorítmicos aplicando estructuras de almacenamiento es el "
        "producto consolidado de esta evidencia de desempeño.",
    )
    add_subheading(document, "Ruta del documento")
    add_table(
        document,
        ("PROBLEMA", "PROPÓSITO", "ARCHIVO FUENTE"),
        tuple((str(problem.number), problem.title, problem.source.name) for problem in PROBLEMS),
        (0.8, 3.6, 2.4),
    )
    add_subheading(document, "Criterios transversales")
    add_bullets(
        document,
        (
            "Cada entrada se valida antes de participar en un cálculo o almacenarse.",
            "La lógica de negocio se separa de la interfaz interactiva para facilitar pruebas.",
            "Los arreglos se procesan sin mutaciones accidentales y con complejidad documentada.",
            "La encuesta musical incorpora persistencia JSON y operaciones de ingreso, consulta, modificación y eliminación.",
        ),
    )
    add_body(
        document,
        "Los listados incluidos a continuación corresponden a las fuentes verificadas por Node.js. "
        "Los ejemplos de prueba emplean datos sintéticos y no contienen información del aprendiz.",
    )

    for problem in PROBLEMS:
        page_break(document)
        add_heading(
            document,
            f"0{problem.number} · Problema {problem.number}",
            problem.title,
            problem.objective,
        )
        add_subheading(document, "Estrategia de solución")
        add_body(document, problem.strategy)
        add_callout(document, "Ejemplo verificable", problem.example, fill=PALE_BLUE)
        add_subheading(document, "Controles de entrada y consistencia")
        add_bullets(document, problem.validations)
        add_callout(document, "Comprobación automática", problem.tests, fill=PALE_BLUE)
        add_subheading(document, f"Código fuente completo · {problem.source.name}")
        add_code(document, problem.source)

    page_break(document)
    add_heading(
        document,
        "05 · Verificación",
        "Matriz final de cumplimiento",
        "Trazabilidad directa entre la lista de chequeo, el código y las pruebas.",
    )
    add_table(
        document,
        ("IND.", "EVIDENCIA", "RESULTADO"),
        (
            ("1", "Problema 1: cuatro figuras, fórmulas y validaciones.", "3 pruebas aprobadas"),
            ("2", "Problema 2: diez edades, categorías y estadísticas.", "2 pruebas aprobadas"),
            ("3", "Problema 3: mezcla lineal de vectores ascendentes.", "2 pruebas aprobadas"),
            ("4", "Problema 4: encuesta, CRUD y persistencia JSON.", "6 pruebas aprobadas"),
        ),
        (0.65, 4.55, 1.6),
    )
    add_callout(
        document,
        "Ejecución verificada",
        "13 de 13 pruebas aprobadas, 0 fallidas. Además, Node.js comprobó la sintaxis de las cuatro fuentes.",
    )
    add_subheading(document, "Conclusiones")
    add_bullets(
        document,
        (
            "Las funciones pequeñas y puras simplifican la lectura, el mantenimiento y las pruebas.",
            "Los arreglos permiten representar colecciones y resolver estadísticas o mezclas de forma eficiente.",
            "Separar dominio, interacción y persistencia reduce errores al administrar información.",
            "La validación previa evita resultados incoherentes y protege el archivo de datos.",
        ),
    )
    add_subheading(document, "Referencias")
    references = (
        "[1] SENA. Guía de aprendizaje GA3-220501093-AA3, páginas 9–10.",
        "[2] MDN Web Docs. Array — https://developer.mozilla.org/docs/Web/JavaScript/Reference/Global_Objects/Array",
        "[3] Node.js. File system — https://nodejs.org/api/fs.html",
        "[4] Node.js. Readline — https://nodejs.org/api/readline.html",
        "[5] Node.js. Test runner — https://nodejs.org/api/test.html",
    )
    for reference in references:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.22)
        paragraph.paragraph_format.first_line_indent = Inches(-0.22)
        paragraph.paragraph_format.space_after = Pt(5)
        add_run(paragraph, reference, size=8.4)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def pdf_text(path: Path) -> str:
    pdftotext = shutil.which("pdftotext")
    if not pdftotext:
        raise RuntimeError("pdftotext es obligatorio para validar la entrega PDF.")
    result = subprocess.run(
        [pdftotext, str(path), "-"],
        check=True,
        capture_output=True,
        text=True,
    )
    return result.stdout


def validate_pdf(
    path: Path,
    profile: dict[str, str] | None,
    *,
    forbidden_values: Iterable[str] = (),
) -> None:
    pdfinfo = shutil.which("pdfinfo")
    if not pdfinfo:
        raise RuntimeError("pdfinfo es obligatorio para validar la entrega PDF.")
    info = subprocess.run(
        [pdfinfo, str(path)],
        check=True,
        capture_output=True,
        text=True,
    ).stdout
    pages = re.search(r"^Pages:\s+(\d+)$", info, flags=re.MULTILINE)
    if not pages or int(pages.group(1)) < 12:
        raise RuntimeError("El PDF no contiene el desarrollo completo esperado.")
    encrypted = re.search(r"^Encrypted:\s+(\S+)", info, flags=re.MULTILINE)
    if encrypted and encrypted.group(1).casefold() != "no":
        raise RuntimeError("El PDF está cifrado y no puede auditarse.")
    expected_author = profile["nombre_completo"] if profile else PUBLIC_AUTHOR
    author = re.search(r"^Author:\s*(.*)$", info, flags=re.MULTILINE)
    if not author or author.group(1).strip() != expected_author:
        raise RuntimeError("El PDF no contiene el autor esperado en sus metadatos.")

    text = pdf_text(path)
    required_terms = (
        "Resolución a problemas algorítmicos aplicando estructuras de almacenamiento",
        "Problema 1",
        "Problema 2",
        "Problema 3",
        "Problema 4",
        "01_figuras_planas.js",
        "02_analisis_edades.js",
        "03_mezclar_vectores.js",
        "04_encuesta_musical.js",
        "13 de 13",
        "Matriz final de cumplimiento",
        "Referencias",
    )
    missing = [term for term in required_terms if term.casefold() not in text.casefold()]
    if missing:
        raise RuntimeError(f"El PDF no contiene secciones obligatorias: {missing}")
    if profile:
        for value in (profile["nombre_completo"], profile["documento"]):
            if value.casefold() not in text.casefold():
                raise RuntimeError("El PDF personalizado no incorporó todos los datos del perfil.")
    for value in forbidden_values:
        if value and value.casefold() in text.casefold():
            raise RuntimeError("El PDF público contiene un valor reservado del perfil local.")


def atomic_publish(source: Path, destination: Path, mode: int) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{destination.name}.",
        suffix=".tmp",
        dir=destination.parent,
    )
    os.close(descriptor)
    temporary = Path(temporary_name)
    try:
        shutil.copyfile(source, temporary)
        os.chmod(temporary, mode)
        os.replace(temporary, destination)
        os.chmod(destination, mode)
    finally:
        temporary.unlink(missing_ok=True)


def create_pdf(
    destination: Path,
    profile: dict[str, str] | None,
    *,
    mode: int,
    forbidden_values: Iterable[str] = (),
) -> None:
    office = shutil.which("libreoffice") or shutil.which("soffice")
    if not office:
        raise RuntimeError("LibreOffice es obligatorio para crear la entrega PDF.")
    with tempfile.TemporaryDirectory(prefix="ev02-documento-") as temp_directory:
        temporary_root = Path(temp_directory)
        docx = temporary_root / f"{CODE}.docx"
        pdf_directory = temporary_root / "pdf"
        pdf_directory.mkdir()
        build_document(docx, profile)
        subprocess.run(
            [office, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_directory), str(docx)],
            check=True,
        )
        generated = pdf_directory / f"{CODE}.pdf"
        if not generated.is_file() or generated.stat().st_size == 0:
            raise RuntimeError("LibreOffice no produjo el PDF esperado.")
        validate_pdf(generated, profile, forbidden_values=forbidden_values)
        atomic_publish(generated, destination, mode)
    validate_pdf(destination, profile, forbidden_values=forbidden_values)


def prepare_exclusive_directory(directory: Path, destination: Path, *, private: bool) -> None:
    if directory.is_symlink() or destination.is_symlink():
        raise RuntimeError("La carpeta de entrega y su archivo no pueden ser enlaces simbólicos.")
    directory.mkdir(parents=True, exist_ok=True)
    if private:
        os.chmod(directory, 0o700)
    unexpected = [path for path in directory.iterdir() if path != destination]
    if unexpected:
        raise RuntimeError(
            f"{directory.name} debe contener un único PDF; archivos inesperados: "
            f"{sorted(path.name for path in unexpected)}"
        )


def ensure_exclusive_file(directory: Path, destination: Path) -> None:
    contents = list(directory.iterdir())
    if contents != [destination] or not destination.is_file() or destination.is_symlink():
        raise RuntimeError(f"{directory.name} no contiene exactamente el PDF esperado.")


def create_public_pdf(profile: dict[str, str] | None) -> Path:
    prepare_exclusive_directory(PUBLIC_DIR, PUBLIC_PDF, private=False)
    forbidden = () if profile is None else (profile["nombre_completo"], profile["documento"])
    create_pdf(PUBLIC_PDF, None, mode=0o644, forbidden_values=forbidden)
    ensure_exclusive_file(PUBLIC_DIR, PUBLIC_PDF)
    return PUBLIC_PDF


def create_personal_pdf(profile: dict[str, str]) -> Path:
    destination = PERSONAL_DIR / (
        f"ENTREGAR_{filename_fragment(profile['nombre_completo'])}_{CODE}.pdf"
    )
    if not git_ignored(destination):
        raise RuntimeError("La entrega personalizada no está protegida por .gitignore.")
    if git_tracked(destination):
        raise RuntimeError("La entrega personalizada aparece registrada en Git.")
    prepare_exclusive_directory(PERSONAL_DIR, destination, private=True)
    create_pdf(destination, profile, mode=0o600)
    os.chmod(PROFILE_PATH, 0o600)
    ensure_exclusive_file(PERSONAL_DIR, destination)
    return destination


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--solo-publico",
        action="store_true",
        help="omite la edición personalizada aunque exista el perfil local",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    validate_sources()
    profile = load_profile()
    public = create_public_pdf(profile)
    print(f"Creado PDF público: {public}")
    if profile is not None and not args.solo_publico:
        personal = create_personal_pdf(profile)
        print(f"Creado un único PDF personalizado local: {personal}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
