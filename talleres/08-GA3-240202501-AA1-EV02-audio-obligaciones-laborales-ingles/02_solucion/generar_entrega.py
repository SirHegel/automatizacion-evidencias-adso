#!/usr/bin/env python3
"""Genera dos PDF públicos y dos PDF personalizados para preparar el audio."""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


SCRIPT_PATH = Path(__file__).resolve()
SOLUTION_DIR = SCRIPT_PATH.parent
WORKSHOP_DIR = SOLUTION_DIR.parent
REPO_ROOT = WORKSHOP_DIR.parents[1]
PUBLIC_DIR = WORKSHOP_DIR / "03_entrega"
PERSONAL_DIR = WORKSHOP_DIR / "04_entrega_personalizada.local"
PROFILE_PATH = REPO_ROOT / "perfil-aprendiz.local.json"

CODE = "GA3-240202501-AA1-EV02"
PUBLIC_AUTHOR = "Entrega académica pública"
PUBLIC_SCRIPT = PUBLIC_DIR / f"{CODE}_Guion_Audio_PUBLICO.pdf"
PUBLIC_PRONUNCIATION = PUBLIC_DIR / f"{CODE}_Guia_Pronunciacion_PUBLICO.pdf"
SCRIPT_SOURCE = SOLUTION_DIR / "GUION_INGLES.md"
PRONUNCIATION_SOURCE = SOLUTION_DIR / "GUION_PRONUNCIACION.md"

GREEN = "39A900"
DARK_GREEN = "174C2C"
DEEP_GREEN = "0D3522"
PALE_GREEN = "EDF7E9"
PALE_BLUE = "EAF3F7"
PALE_GOLD = "FFF4D6"
PALE_RED = "FCECE8"
INK = "17231B"
MUTED = "5A685F"
LINE = "D7DFD9"
WHITE = "FFFFFF"

PUBLIC_VALUES = {
    "name": "[LEARNER NAME]",
    "document": "[DOCUMENT ID]",
    "group": "[TRAINING GROUP]",
    "instructor": "[INSTRUCTOR]",
    "audio_url": "[PUBLIC AUDIO LINK PENDING]",
}


def run(command: list[str], *, capture: bool = False) -> subprocess.CompletedProcess[str]:
    printable = " ".join(command)
    print(f"→ {printable}")
    return subprocess.run(
        command,
        cwd=REPO_ROOT,
        check=True,
        text=True,
        capture_output=capture,
    )


def extract_block(path: Path, start: str, end: str) -> str:
    text = path.read_text(encoding="utf-8")
    before, separator, remainder = text.partition(start)
    if not separator:
        raise RuntimeError(f"No se encontró {start!r} en {path}.")
    block, separator, after = remainder.partition(end)
    if not separator or not before or not after:
        raise RuntimeError(f"El bloque {start!r} de {path} está incompleto.")
    return block.strip()


def canonical_script() -> list[str]:
    block = extract_block(SCRIPT_SOURCE, "<!-- SCRIPT_START -->", "<!-- SCRIPT_END -->")
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", block) if part.strip()]
    if len(paragraphs) != 7:
        raise RuntimeError(f"El guion debe contener siete párrafos; contiene {len(paragraphs)}.")
    text = "\n\n".join(paragraphs)
    count = len(re.findall(r"[A-Za-z]+(?:['’][A-Za-z]+)?|\d+", text))
    if count != 386:
        raise RuntimeError(f"El guion público cambió de extensión: {count} palabras, se esperaban 386.")
    expected_counts = {
        "have/has to": len(re.findall(r"\b(?:have|has) to\b", text, flags=re.IGNORECASE)),
        "must": len(re.findall(r"\bmust\b", text, flags=re.IGNORECASE)),
        "should": len(re.findall(r"\bshould\b", text, flags=re.IGNORECASE)),
    }
    if expected_counts != {"have/has to": 7, "must": 8, "should": 9}:
        raise RuntimeError(f"Conteo modal inesperado: {expected_counts}.")
    required = (
        "This recording is for the activity Audio GA3-240202501-AA1-EV02.",
        "In my opinion",
        "I believe",
        "I think",
        "academic",
        "software development",
        "deadline",
        "error in the application",
    )
    missing = [term for term in required if term.casefold() not in text.casefold()]
    if missing:
        raise RuntimeError(f"El guion perdió contenidos obligatorios: {missing}.")
    return paragraphs


def pronunciation_rows() -> list[tuple[str, str]]:
    block = extract_block(
        PRONUNCIATION_SOURCE,
        "<!-- PRONUNCIATION_START -->",
        "<!-- PRONUNCIATION_END -->",
    )
    rows: list[tuple[str, str]] = []
    for line in block.splitlines():
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if len(cells) != 2 or cells[0] in {"English segment", "---"}:
            continue
        if set(cells[0]) == {"-"}:
            continue
        rows.append((cells[0], cells[1]))
    sentences = [
        sentence.strip()
        for paragraph in canonical_script()
        for sentence in re.split(r"(?<=[.!?])\s+", paragraph)
        if sentence.strip()
    ]
    english_rows = [english for english, support in rows]
    if english_rows != sentences:
        for position, (expected, actual) in enumerate(
            zip(sentences, english_rows, strict=False),
            start=1,
        ):
            if expected != actual:
                raise RuntimeError(
                    "La pronunciación dejó de corresponder con el guion en el segmento "
                    f"{position}: {expected!r} != {actual!r}."
                )
        raise RuntimeError(
            "La pronunciación y el guion tienen distinta cantidad de segmentos: "
            f"{len(rows)} frente a {len(sentences)}."
        )
    code_support = next(
        support for english, support in rows if english.startswith("This recording is")
    )
    expected_digits = "tu FOR ZI-rou tu ZI-rou tu FAIV ZI-rou WAN"
    if expected_digits not in code_support:
        raise RuntimeError("La guía de pronunciación no lee correctamente el código numérico.")
    return rows


def filename_fragment(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii")
    fragment = re.sub(r"[^A-Za-z0-9]+", "_", ascii_value).strip("_")
    if not fragment:
        raise RuntimeError("El nombre local no permite construir un archivo seguro.")
    return fragment


def load_profile() -> dict[str, str] | None:
    if not PROFILE_PATH.is_file():
        return None
    profile = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    required = ("nombre_completo", "tipo_documento", "documento", "programa", "institucion")
    missing = [key for key in required if not str(profile.get(key, "")).strip()]
    if missing:
        raise RuntimeError(f"El perfil local no contiene campos obligatorios: {missing}.")
    return {key: str(value).strip() for key, value in profile.items()}


def personalized_values(profile: dict[str, str]) -> tuple[dict[str, str], list[str]]:
    optional = {
        "group": ("grupo_formacion", "[PENDING: TRAINING GROUP]"),
        "instructor": ("instructor", "[PENDING: INSTRUCTOR NAME]"),
        "audio_url": ("enlace_audio", "[PENDING: PUBLIC AUDIO LINK]"),
    }
    values = {
        "name": profile["nombre_completo"],
        "document": profile["documento"],
    }
    missing: list[str] = []
    for target, (source, placeholder) in optional.items():
        value = profile.get(source, "").strip()
        values[target] = value or placeholder
        if not value:
            missing.append(target)
    return values, missing


def replace_placeholders(text: str, values: dict[str, str]) -> str:
    replacements = {
        "[LEARNER NAME]": values["name"],
        "[DOCUMENT ID]": values["document"],
        "[TRAINING GROUP]": values["group"],
        "[INSTRUCTOR]": values["instructor"],
    }
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, value)
    return text


def digit_pronunciation(value: str) -> str:
    words = {
        "0": "zero",
        "1": "one",
        "2": "two",
        "3": "three",
        "4": "four",
        "5": "five",
        "6": "six",
        "7": "seven",
        "8": "eight",
        "9": "nine",
    }
    compact = re.sub(r"[\s.-]", "", value)
    if compact.isdigit():
        return " ".join(words[digit] for digit in compact)
    return "[say each digit slowly]"


def set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, *, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    value = OxmlElement("w:t")
    value.text = text
    run.extend((properties, value))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(f"{CODE}  ·  ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def configure_document(document: Document, *, author: str, title: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for style_name, size, color in (
        ("Title", 26, DEEP_GREEN),
        ("Heading 1", 18, DARK_GREEN),
        ("Heading 2", 13, DARK_GREEN),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    properties = document.core_properties
    properties.author = author
    properties.last_modified_by = author
    properties.title = title
    properties.subject = CODE
    properties.keywords = "English audio, obligations, pronunciation, SENA"
    properties.comments = "Generated reproducibly from public sources."

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("AUDIO · ENGLISH · A2.2")
    header_run.bold = True
    header_run.font.name = "Arial"
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor.from_string(GREEN)
    add_page_number(section.footer.paragraphs[0])


def add_label(document: Document, text: str, *, color: str = PALE_GREEN, ink: str = DARK_GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(7.0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, top=120, bottom=120)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(ink)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_cover(
    document: Document,
    *,
    title: str,
    subtitle: str,
    edition_label: str,
    edition_color: str,
    edition_ink: str,
    values: dict[str, str],
    include_audio_link: bool,
) -> None:
    accent = document.add_table(rows=1, cols=1)
    accent.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_shading(accent.cell(0, 0), GREEN)
    accent.cell(0, 0).height = Inches(0.14)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(22)
    code = document.add_paragraph()
    code.alignment = WD_ALIGN_PARAGRAPH.CENTER
    code_run = code.add_run(CODE)
    code_run.bold = True
    code_run.font.name = "Arial"
    code_run.font.size = Pt(11)
    code_run.font.color.rgb = RGBColor.from_string(GREEN)

    heading = document.add_paragraph(style="Title")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run(title)
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(subtitle)
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor.from_string(MUTED)
    sub.paragraph_format.space_after = Pt(18)

    add_label(document, edition_label, color=edition_color, ink=edition_ink)

    fields = [
        ("Learner", values["name"]),
        ("Identification", values["document"]),
        ("Training group", values["group"]),
        ("Instructor", values["instructor"]),
        ("Program", "Software Analysis and Development"),
        ("Activity", f"Audio {CODE}"),
    ]
    if include_audio_link:
        fields.append(("Public audio", values["audio_url"]))
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (label, value) in enumerate(fields):
        cells = table.add_row().cells
        cells[0].width = Inches(1.55)
        cells[1].width = Inches(5.4)
        set_cell_shading(cells[0], PALE_GREEN if index % 2 == 0 else PALE_BLUE)
        set_cell_shading(cells[1], "F8FAF8" if index % 2 == 0 else "F5F9FB")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        left = cells[0].paragraphs[0]
        left_run = left.add_run(label)
        left_run.bold = True
        left_run.font.name = "Arial"
        left_run.font.size = Pt(9.5)
        left_run.font.color.rgb = RGBColor.from_string(DARK_GREEN)
        right = cells[1].paragraphs[0]
        if label == "Public audio" and value.casefold().startswith(("https://", "http://")):
            add_hyperlink(right, value, value)
        else:
            right_run = right.add_run(value)
            right_run.font.name = "Arial"
            right_run.font.size = Pt(9.5)
            if value.startswith("[PENDING") or value.startswith("[PUBLIC"):
                right_run.bold = True
                right_run.font.color.rgb = RGBColor.from_string("9B4B20")

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(16)
    note_run = note.add_run("Prepared for a 2–5 minute recording with the learner's real voice.")
    note_run.italic = True
    note_run.font.name = "Arial"
    note_run.font.size = Pt(9.5)
    note_run.font.color.rgb = RGBColor.from_string(MUTED)
    document.add_page_break()


def add_highlighted_text(paragraph, text: str) -> None:
    pattern = re.compile(
        r"\b(?:have to|has to|must(?: not)?|should)\b|"
        r"\b(?:I believe|In my opinion|I think)\b",
        flags=re.IGNORECASE,
    )
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        run = paragraph.add_run(match.group(0))
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(GREEN)
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)


def add_script_body(document: Document, paragraphs: list[str], values: dict[str, str]) -> None:
    document.add_heading("Recording script", level=1)
    intro = document.add_paragraph(
        "Read only the English text below. Green bold expressions show the required "
        "modal structures and opinion phrases; do not say the headings aloud."
    )
    intro.paragraph_format.space_after = Pt(10)

    section_titles = (
        "1 · Introduction",
        "2 · Academic responsibilities",
        "3 · Workplace responsibilities",
        "4 · Solution to a missed deadline",
        "5 · Solution to an application error",
        "6 · Attitudes and beliefs",
        "7 · Conclusion",
    )
    for title, source in zip(section_titles, paragraphs):
        heading = document.add_paragraph()
        set_keep_with_next(heading)
        heading.paragraph_format.space_before = Pt(5)
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(DARK_GREEN)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.1
        paragraph.paragraph_format.space_after = Pt(6)
        add_highlighted_text(paragraph, replace_placeholders(source, values))

    document.add_heading("Final recording checklist", level=1)
    checklist = (
        "Say the personal introduction slowly and keep a short pause between sections.",
        "Stress have to, must and should without adding -ing after them.",
        "Use calm falling intonation, then listen and confirm a duration of two to five minutes.",
        "If a link is submitted, test it in a private browser window before delivery.",
    )
    for item in checklist:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_pronunciation_body(
    document: Document,
    rows: list[tuple[str, str]],
    values: dict[str, str],
) -> None:
    document.add_heading("Pronunciation and delivery guide", level=1)
    paragraph = document.add_paragraph(
        "Use the English column as the only text to record. The second column is a "
        "practical approximation for rehearsal; it is not a translation."
    )
    paragraph.paragraph_format.space_after = Pt(8)

    key = document.add_table(rows=2, cols=4)
    key.alignment = WD_TABLE_ALIGNMENT.CENTER
    key.autofit = False
    key_values = (
        ("CAPITALS", "strong syllable", "/", "short pause"),
        ("//", "longer pause", "th", "tongue lightly between the teeth"),
    )
    for row, values_row in zip(key.rows, key_values):
        for index, (cell, value) in enumerate(zip(row.cells, values_row)):
            set_cell_shading(cell, PALE_GREEN if index % 2 == 0 else "F8FAF8")
            set_cell_margins(cell, top=70, bottom=70)
            run = cell.paragraphs[0].add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.bold = index % 2 == 0
            run.font.color.rgb = RGBColor.from_string(DARK_GREEN if index % 2 == 0 else INK)

    document.add_paragraph()
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)
    headers = table.rows[0].cells
    for cell, label in zip(headers, ("English segment · RECORD THIS", "Practice support · DO NOT RECORD")):
        set_cell_shading(cell, DARK_GREEN)
        set_cell_margins(cell, top=95, bottom=95)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    set_repeat_table_header(table.rows[0])

    for index, (english, support) in enumerate(rows, start=1):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=85, bottom=85, start=95, end=95)
        set_cell_shading(cells[0], "F8FAF8" if index % 2 else PALE_GREEN)
        set_cell_shading(cells[1], "FFFFFF" if index % 2 else PALE_BLUE)
        english_text = replace_placeholders(english, values)
        support_text = support
        support_text = support_text.replace("[say your full name clearly]", values["name"])
        support_text = support_text.replace(
            "[say each digit]",
            digit_pronunciation(values["document"]),
            1,
        )
        support_text = support_text.replace(
            "[say each digit]",
            digit_pronunciation(values["group"]),
            1,
        )
        support_text = support_text.replace(
            "[say your instructor's name]",
            values["instructor"],
        )
        for cell, value, color in (
            (cells[0], english_text, INK),
            (cells[1], support_text, MUTED),
        ):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.keep_together = True
            run = paragraph.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(8.3)
            run.font.color.rgb = RGBColor.from_string(color)

    document.add_heading("Practice sequence", level=1)
    for item in (
        "Practice one row at a time, then read the complete paragraph without stopping.",
        "Keep final consonants in must, should, work and respect.",
        "Do not pronounce the letter l in should.",
        "Use a falling tone at the end of statements and a gentle continuing tone inside lists.",
        "Record a short test before making the complete audio.",
    ):
        document.add_paragraph(item, style="List Number")


def build_script_document(
    *,
    values: dict[str, str],
    author: str,
    label: str,
    label_color: str,
    label_ink: str,
) -> Document:
    document = Document()
    configure_document(
        document,
        author=author,
        title=f"{CODE} — English audio script",
    )
    add_cover(
        document,
        title="English audio script",
        subtitle="Academic and workplace responsibilities in software development",
        edition_label=label,
        edition_color=label_color,
        edition_ink=label_ink,
        values=values,
        include_audio_link=True,
    )
    add_script_body(document, canonical_script(), values)
    return document


def build_pronunciation_document(
    *,
    values: dict[str, str],
    author: str,
    label: str,
    label_color: str,
    label_ink: str,
) -> Document:
    document = Document()
    configure_document(
        document,
        author=author,
        title=f"{CODE} — Pronunciation guide",
    )
    add_cover(
        document,
        title="Pronunciation guide",
        subtitle="Sentence-by-sentence support for the English recording",
        edition_label=label,
        edition_color=label_color,
        edition_ink=label_ink,
        values=values,
        include_audio_link=False,
    )
    add_pronunciation_body(document, pronunciation_rows(), values)
    return document


def export_document(document: Document, destination: Path, *, mode: int) -> None:
    office = shutil.which("libreoffice") or shutil.which("soffice")
    if not office:
        raise RuntimeError("LibreOffice no está disponible para exportar los PDF.")
    destination.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="audio-ev02-") as temporary_name:
        temporary = Path(temporary_name)
        source = temporary / f"{destination.stem}.docx"
        document.save(source)
        run(
            [
                office,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(temporary),
                str(source),
            ]
        )
        generated = temporary / f"{source.stem}.pdf"
        if not generated.is_file() or generated.stat().st_size == 0:
            raise RuntimeError(f"LibreOffice no creó el PDF esperado: {generated}.")
        descriptor, staging_name = tempfile.mkstemp(
            prefix=f".{destination.name}.",
            suffix=".tmp",
            dir=destination.parent,
        )
        os.close(descriptor)
        staging = Path(staging_name)
        try:
            shutil.copyfile(generated, staging)
            os.chmod(staging, mode)
            os.replace(staging, destination)
            os.chmod(destination, mode)
        finally:
            if staging.exists():
                staging.unlink()


def git_ignored(path: Path) -> bool:
    relative = path.relative_to(REPO_ROOT)
    result = subprocess.run(
        ["git", "check-ignore", "--no-index", "-q", "--", relative.as_posix()],
        cwd=REPO_ROOT,
        check=False,
    )
    return result.returncode == 0


def git_tracked(path: Path) -> bool:
    relative = path.relative_to(REPO_ROOT)
    result = subprocess.run(
        ["git", "ls-files", "--error-unmatch", "--", relative.as_posix()],
        cwd=REPO_ROOT,
        check=False,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return result.returncode == 0


def validate_personal_destination(destination: Path) -> None:
    if PERSONAL_DIR.exists() and PERSONAL_DIR.is_symlink():
        raise RuntimeError("La carpeta personalizada no puede ser un enlace simbólico.")
    if destination.parent != PERSONAL_DIR:
        raise RuntimeError("La entrega personalizada salió de la carpeta autorizada.")
    if not git_ignored(destination):
        raise RuntimeError(f"Git no está ignorando la entrega personalizada: {destination}.")
    if git_tracked(destination):
        raise RuntimeError(f"La entrega personalizada ya está registrada en Git: {destination}.")


def pdf_text(path: Path) -> str:
    pdftotext = shutil.which("pdftotext")
    if not pdftotext:
        raise RuntimeError("pdftotext no está disponible para validar los entregables.")
    return run([pdftotext, str(path), "-"], capture=True).stdout


def validate_pdf(path: Path, *, expected_author: str, required_terms: tuple[str, ...]) -> str:
    pdfinfo = shutil.which("pdfinfo")
    if not pdfinfo:
        raise RuntimeError("pdfinfo no está disponible para validar los entregables.")
    information = run([pdfinfo, str(path)], capture=True).stdout
    pages = re.search(r"^Pages:\s+(\d+)$", information, flags=re.MULTILINE)
    author = re.search(r"^Author:\s*(.*)$", information, flags=re.MULTILINE)
    encrypted = re.search(r"^Encrypted:\s*(\S+)", information, flags=re.MULTILINE)
    if not pages or int(pages.group(1)) < 2:
        raise RuntimeError(f"El PDF no tiene una paginación válida: {path}.")
    if not author or author.group(1).strip() != expected_author:
        raise RuntimeError(
            f"Autor inesperado en {path}: {author.group(1).strip() if author else 'ausente'}.")
    if encrypted and encrypted.group(1).casefold() != "no":
        raise RuntimeError(f"El PDF está cifrado: {path}.")
    text = pdf_text(path)
    missing = [term for term in required_terms if term.casefold() not in text.casefold()]
    if missing:
        raise RuntimeError(f"El PDF {path.name} no contiene {missing}.")
    return text


def generate_public() -> tuple[Path, Path]:
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
    expected = {PUBLIC_SCRIPT, PUBLIC_PRONUNCIATION}
    unexpected = {path for path in PUBLIC_DIR.iterdir() if path.is_file()} - expected
    if unexpected:
        raise RuntimeError(
            "03_entrega debe contener solo los dos PDF públicos del taller 8; sobran: "
            f"{sorted(path.name for path in unexpected)}."
        )
    export_document(
        build_script_document(
            values=PUBLIC_VALUES,
            author=PUBLIC_AUTHOR,
            label="PUBLIC VERSION · PERSONAL DATA REMOVED",
            label_color=PALE_BLUE,
            label_ink=DARK_GREEN,
        ),
        PUBLIC_SCRIPT,
        mode=0o644,
    )
    export_document(
        build_pronunciation_document(
            values=PUBLIC_VALUES,
            author=PUBLIC_AUTHOR,
            label="PUBLIC VERSION · PRONUNCIATION SUPPORT",
            label_color=PALE_BLUE,
            label_ink=DARK_GREEN,
        ),
        PUBLIC_PRONUNCIATION,
        mode=0o644,
    )
    if {path for path in PUBLIC_DIR.iterdir() if path.is_file()} != expected:
        raise RuntimeError("03_entrega no contiene exactamente los dos PDF públicos.")
    script_text = validate_pdf(
        PUBLIC_SCRIPT,
        expected_author=PUBLIC_AUTHOR,
        required_terms=(
            f"Audio {CODE}",
            "have to",
            "must",
            "should",
            "misses a deadline",
            "error in the application",
        ),
    )
    pronunciation_text = validate_pdf(
        PUBLIC_PRONUNCIATION,
        expected_author=PUBLIC_AUTHOR,
        required_terms=(
            f"Audio {CODE}",
            "Pronunciation guide",
            "have to",
            "must",
            "should",
            "DO NOT RECORD",
        ),
    )
    for placeholder in PUBLIC_VALUES.values():
        if placeholder not in script_text and placeholder not in pronunciation_text:
            raise RuntimeError(f"Los PDF públicos perdieron el marcador {placeholder!r}.")
    return PUBLIC_SCRIPT, PUBLIC_PRONUNCIATION


def generate_personal(profile: dict[str, str]) -> tuple[Path, Path]:
    values, missing = personalized_values(profile)
    fragment = filename_fragment(profile["nombre_completo"])
    final = PERSONAL_DIR / f"ENTREGAR_{fragment}_{CODE}_Guion_y_Enlace.pdf"
    support = PERSONAL_DIR / f"APOYO_GRABACION_{fragment}_{CODE}_Pronunciacion.pdf"
    for destination in (final, support):
        validate_personal_destination(destination)
    PERSONAL_DIR.mkdir(parents=True, exist_ok=True, mode=0o700)
    os.chmod(PERSONAL_DIR, 0o700)
    expected = {final, support}
    unexpected = {path for path in PERSONAL_DIR.iterdir()} - expected
    if unexpected:
        raise RuntimeError(
            "La carpeta personalizada debe contener solo los dos PDF definidos; sobran: "
            f"{sorted(path.name for path in unexpected)}."
        )
    if missing:
        labels = {
            "group": "GRUPO DE FORMACIÓN",
            "instructor": "INSTRUCTOR",
            "audio_url": "ENLACE PÚBLICO DEL AUDIO",
        }
        readable = ", ".join(labels[value] for value in missing)
        final_label = f"BORRADOR · NO LISTO PARA ENTREGAR · FALTAN: {readable.upper()}"
    else:
        final_label = "ENTREGA FINAL · GUION Y ENLACE PÚBLICO · LISTO PARA ENTREGAR"
    export_document(
        build_script_document(
            values=values,
            author=profile["nombre_completo"],
            label=final_label,
            label_color=PALE_RED if missing else PALE_GREEN,
            label_ink="9B2F20" if missing else DARK_GREEN,
        ),
        final,
        mode=0o600,
    )
    export_document(
        build_pronunciation_document(
            values=values,
            author=profile["nombre_completo"],
            label="NO SUBIR · SOLO APOYO PARA GRABAR",
            label_color=PALE_GOLD,
            label_ink="8A4B00",
        ),
        support,
        mode=0o600,
    )
    if {path for path in PERSONAL_DIR.iterdir()} != expected:
        raise RuntimeError("La carpeta personalizada no contiene exactamente los dos PDF esperados.")
    final_text = validate_pdf(
        final,
        expected_author=profile["nombre_completo"],
        required_terms=(profile["nombre_completo"], profile["documento"], f"Audio {CODE}"),
    )
    support_text = validate_pdf(
        support,
        expected_author=profile["nombre_completo"],
        required_terms=(
            profile["nombre_completo"],
            profile["documento"],
            "NO SUBIR",
            "Pronunciation guide",
        ),
    )
    if missing and "NO LISTO PARA ENTREGAR" not in final_text:
        raise RuntimeError("La entrega incompleta no quedó marcada como borrador.")
    if not missing and any("[PENDING:" in text for text in (final_text, support_text)):
        raise RuntimeError("La entrega final aún conserva campos pendientes.")
    return final, support


def validate_public_privacy(profile: dict[str, str] | None) -> None:
    if profile is None:
        return
    public_text = "\n".join((pdf_text(PUBLIC_SCRIPT), pdf_text(PUBLIC_PRONUNCIATION))).casefold()
    sensitive = [profile["nombre_completo"], profile["documento"]]
    sensitive.extend(
        profile[key]
        for key in ("grupo_formacion", "instructor", "enlace_audio")
        if profile.get(key, "").strip()
    )
    leaked = [value for value in sensitive if value.casefold() in public_text]
    if leaked:
        raise RuntimeError("Los PDF públicos contienen valores del perfil local.")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--solo-publico",
        action="store_true",
        help="genera únicamente los dos PDF públicos aunque exista el perfil local",
    )
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    canonical_script()
    pronunciation_rows()
    public_outputs = generate_public()
    profile = None if args.solo_publico else load_profile()
    validate_public_privacy(profile)
    print("PDF públicos creados:")
    for path in public_outputs:
        print(f"  - {path}")
    if profile is not None:
        personal_outputs = generate_personal(profile)
        print("PDF personalizados locales creados:")
        for path in personal_outputs:
            print(f"  - {path}")
    else:
        print("Perfil local ausente u omitido; no se generaron documentos personalizados.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
